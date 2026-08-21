import os, re, time, shutil, stat, io, csv
from fastapi import FastAPI, Request, UploadFile, File
from fastapi.responses import HTMLResponse, StreamingResponse, JSONResponse, Response
from fastapi.staticfiles import StaticFiles
from openai import OpenAI
import json
import asyncio
import requests
from requests.auth import HTTPBasicAuth
from docx import Document
from docx.shared import Pt, Inches
from pypdf import PdfReader
from dotenv import load_dotenv
import db

load_dotenv()

app = FastAPI()
db.init_db()

# ============================================================
# CONFIGURATION
# ============================================================

# Shares db.DATA_DIR (rather than reading DATA_DIR separately) so the DB and
# the generated-files folder always live on the same disk - see db.py for
# why this matters on Render.
WORKSPACE_DIR = os.path.join(db.DATA_DIR, "workspace")
if not os.path.exists(WORKSPACE_DIR):
    os.makedirs(WORKSPACE_DIR)

# Mount Workspace for Live Preview
app.mount("/workspace", StaticFiles(directory=WORKSPACE_DIR), name="workspace")

def get_project_dir(project_id):
    """
    Each app-level project gets its own subfolder under WORKSPACE_DIR so
    generating documents for one project can never overwrite another's.
    Projects created before this isolation existed have their files sitting
    directly in WORKSPACE_DIR root instead - callers fall back to that root
    when a project's own subfolder has nothing in it yet.
    """
    safe_id = re.sub(r'[^A-Za-z0-9_-]', '_', str(project_id or "default"))[:100] or "default"
    project_path = os.path.join(WORKSPACE_DIR, safe_id)
    os.makedirs(project_path, exist_ok=True)
    return project_path, safe_id

# OpenAI - migrated off Gemini after repeatedly hitting its free-tier daily
# request quota (see llm() below for the actual call).
openai_client = OpenAI(api_key=os.getenv("OPENAI_API_KEY", ""))

# ============================================================
# PROJECT PERSISTENCE API (SQLite)
# ============================================================

@app.post("/api/projects")
async def api_create_project(request: Request):
    body = await request.json()
    name = (body.get("name") or "").strip()
    key = (body.get("key") or "").strip().upper()
    description = (body.get("description") or "").strip()
    jira_project_key = (body.get("jira_project_key") or "").strip().upper()
    confluence_space_key = (body.get("confluence_space_key") or "").strip().upper()
    if not name or not key:
        return JSONResponse({"error": "Project name and key are required"}, status_code=400)
    project = db.create_project(name, key, description or "No description provided.", jira_project_key, confluence_space_key)
    return JSONResponse({"success": True, "project": project})


@app.get("/api/projects")
async def api_list_projects():
    return JSONResponse({"success": True, "projects": db.list_projects()})


@app.get("/api/projects/{project_id}")
async def api_get_project(project_id: str):
    project = db.get_project(project_id)
    if not project:
        return JSONResponse({"error": "Project not found"}, status_code=404)
    return JSONResponse({"success": True, "project": project})


def _force_rmtree(path):
    """
    Removes a directory tree, clearing read-only attributes on the way if a
    delete fails. On this machine the workspace folder lives inside a live
    OneDrive sync root, which can mark just-written files read-only while it
    syncs - shutil.rmtree's default behavior raises PermissionError on those
    instead of clearing the attribute like Explorer/PowerShell do.
    """
    def _on_error(func, target_path, exc_info):
        try:
            os.chmod(target_path, stat.S_IWRITE)
            func(target_path)
        except Exception:
            pass
    shutil.rmtree(path, onerror=_on_error)


@app.delete("/api/projects/{project_id}")
async def api_delete_project(project_id: str):
    db.delete_project(project_id)
    project_dir, _ = get_project_dir(project_id)
    if os.path.isdir(project_dir):
        _force_rmtree(project_dir)
    return JSONResponse({"success": True})


@app.post("/api/projects/{project_id}/knowledge")
async def api_add_knowledge_item(project_id: str, request: Request):
    body = await request.json()
    title = (body.get("title") or "").strip()
    content = (body.get("content") or "").strip()
    if not title or not content:
        return JSONResponse({"error": "Title and content are required"}, status_code=400)
    item_id = db.add_knowledge_item(project_id, title, content, body.get("type", "Manual Instruction"))
    return JSONResponse({"success": True, "id": item_id})


@app.delete("/api/projects/{project_id}/knowledge/{item_id}")
async def api_delete_knowledge_item(project_id: str, item_id: int):
    db.delete_knowledge_item(item_id)
    return JSONResponse({"success": True})


KNOWLEDGE_UPLOAD_MAX_CHARS = 50000

@app.post("/api/projects/{project_id}/knowledge/upload")
async def api_upload_knowledge_file(project_id: str, file: UploadFile = File(...)):
    """
    Extracts text from an uploaded reference document (.txt/.md/.docx/.pdf)
    and stores it as a knowledge item, same as a manually pasted note - it
    gets included as context in every subsequent agent prompt for this project.
    """
    filename = file.filename or "Uploaded Document"
    ext = os.path.splitext(filename)[1].lower()
    raw = await file.read()

    try:
        if ext in (".txt", ".md"):
            text = raw.decode("utf-8", errors="ignore")
        elif ext == ".docx":
            text = "\n".join(p.text for p in Document(io.BytesIO(raw)).paragraphs)
        elif ext == ".pdf":
            reader = PdfReader(io.BytesIO(raw))
            text = "\n".join((page.extract_text() or "") for page in reader.pages)
        else:
            return JSONResponse({"success": False, "error": f"Unsupported file type '{ext}'. Supported: .txt, .md, .docx, .pdf"}, status_code=400)
    except Exception as e:
        return JSONResponse({"success": False, "error": f"Failed to read file: {str(e)}"}, status_code=400)

    text = text.strip()
    if not text:
        return JSONResponse({"success": False, "error": "No extractable text found in this file."}, status_code=400)

    truncated = len(text) > KNOWLEDGE_UPLOAD_MAX_CHARS
    if truncated:
        text = text[:KNOWLEDGE_UPLOAD_MAX_CHARS] + "\n\n[...truncated...]"

    title = os.path.splitext(filename)[0]
    item_id = db.add_knowledge_item(project_id, title, text, "Uploaded Document")
    return JSONResponse({"success": True, "item": {"id": item_id, "title": title, "content": text, "type": "Uploaded Document"}, "truncated": truncated})

# ============================================================
# JIRA CONFIGURATION
# ============================================================

JIRA_URL          = os.getenv("JIRA_URL", "").rstrip("/")
JIRA_USERNAME     = os.getenv("JIRA_USERNAME", "")
JIRA_API_TOKEN    = os.getenv("JIRA_API_TOKEN", "")
JIRA_PERSONAL_TOKEN = os.getenv("JIRA_PERSONAL_TOKEN", "")
JIRA_PROJECT_KEY  = os.getenv("JIRA_PROJECT_KEY", "")
JIRA_SSL_VERIFY   = os.getenv("JIRA_SSL_VERIFY", "true").lower() != "false"

def jira_headers():
    if JIRA_PERSONAL_TOKEN:
        return {
            "Authorization": f"Bearer {JIRA_PERSONAL_TOKEN}",
            "Content-Type": "application/json",
            "Accept": "application/json",
        }, None
    else:
        return {
            "Content-Type": "application/json",
            "Accept": "application/json",
        }, HTTPBasicAuth(JIRA_USERNAME, JIRA_API_TOKEN)

def jira_configured():
    if not JIRA_URL or not JIRA_PROJECT_KEY:
        return False
    if JIRA_PERSONAL_TOKEN:
        return True
    return bool(JIRA_USERNAME and JIRA_API_TOKEN)

def resolve_project_keys(project_id):
    """
    Returns (jira_key, confluence_key) for a project. Falls back to the
    global .env JIRA_PROJECT_KEY when the project didn't set its own. If a
    project set its own Jira key but not a separate Confluence space key,
    the Jira key is reused for Confluence too (matches the common Atlassian
    convention of one key shared across both products on a site).
    """
    jira_key = JIRA_PROJECT_KEY
    confluence_key = JIRA_PROJECT_KEY
    if project_id:
        proj = db.get_project_basic(project_id)
        if proj:
            if proj.get("jira_project_key"):
                jira_key = proj["jira_project_key"]
                confluence_key = proj["jira_project_key"]
            if proj.get("confluence_space_key"):
                confluence_key = proj["confluence_space_key"]
    return jira_key, confluence_key


def _sanitize_work_item_tree(nodes, parent_type=None):
    """
    Enforces this Jira site's real 3-level hierarchy on LLM-generated JSON
    (verified via /rest/api/3/issuetype: Epic=1, Story=0, Task=0, Subtask=-1):
    Epic at the top; Story and Task are peers, both valid children of an Epic
    or standalone top-level items; Subtask is the only level below, valid
    only under a Story or a Task. Nodes that don't fit their position are
    dropped rather than silently producing a tree Jira would reject on push.
    """
    if not isinstance(nodes, list):
        return []

    if parent_type is None:
        allowed = {"Epic", "Story", "Task"}
    elif parent_type == "Epic":
        allowed = {"Story", "Task"}
    elif parent_type in ("Story", "Task"):
        allowed = {"Subtask"}
    else:
        allowed = set()  # Subtask can't have children

    cleaned = []
    for node in nodes:
        if not isinstance(node, dict):
            continue
        issuetype = node.get("issuetype")
        # "Task" is never valid directly under a Story/Task (Task is always a
        # peer of Story, never nested under one) - but the LLM sometimes still
        # emits it there when the request said "task" colloquially (e.g. "add
        # 3 tasks under this story"). Coerce rather than drop: Subtask is the
        # one type that actually fits, and it's clearly what was meant.
        if issuetype == "Task" and allowed == {"Subtask"}:
            issuetype = "Subtask"
        if issuetype not in allowed:
            continue
        summary = (node.get("summary") or "").strip()
        if not summary:
            continue
        cleaned.append({
            "issuetype": issuetype,
            "summary": summary,
            "description": node.get("description") or "",
            "priority": node.get("priority") or "Medium",
            "children": _sanitize_work_item_tree(node.get("children"), issuetype),
        })
    return cleaned

def get_confluence_space_id(space_key, headers, auth):
    """Resolves a Confluence space key to the numeric space ID the v2 API requires."""
    get_kwargs = {"headers": headers, "params": {"keys": space_key}, "verify": JIRA_SSL_VERIFY, "timeout": 15}
    if auth:
        get_kwargs["auth"] = auth
    space_r = requests.get(f"{JIRA_URL}/wiki/api/v2/spaces", **get_kwargs)
    space_results = space_r.json().get("results", []) if space_r.status_code == 200 else []
    return space_results[0]["id"] if space_results else None


# ============================================================
# LLM & DOCX UTILITY
# ============================================================

def llm(system, user):
    try:
        model_name = os.getenv("OPENAI_MODEL", "gpt-4o-mini")
        response = openai_client.chat.completions.create(
            model=model_name,
            messages=[
                {"role": "system", "content": system + "\n\nImportant: Do not recite copyrighted material or training data verbatim. Be creative and synthesize new logic based on requirements."},
                {"role": "user", "content": user},
            ],
            temperature=0.7,
        )
        choice = response.choices[0]
        if choice.finish_reason == "content_filter":
            return "AGENT ERROR: Content filter block. Try re-phrasing your request slightly."
        return (choice.message.content or "").strip()
    except Exception as e:
        return f"LLM_ERROR: {str(e)}"


def _is_llm_error(text):
    """llm() returns its failure (quota exceeded, recitation block, API
    error, etc.) as a normal string rather than raising, so every call site
    that persists an llm() result as document content MUST check this before
    writing/overwriting a file or DB row - otherwise a transient API failure
    silently becomes the saved content, clobbering whatever real content was
    there before with no way to recover it."""
    return isinstance(text, str) and (text.startswith("LLM_ERROR:") or text.startswith("AGENT ERROR:"))

def save_as_docx(title, content, file_path):
    """
    Parses Markdown-like headers and paragraphs to output a clean DOCX file.
    """
    doc = Document()
    
    # Configure styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    doc.add_heading(title, level=0)
    
    lines = content.split('\n')
    for line in lines:
        line_strip = line.strip()
        if not line_strip:
            continue
        
        # Headers
        if line_strip.startswith('###'):
            h = doc.add_heading(line_strip.replace('###', '').strip(), level=3)
            h.paragraph_format.space_before = Pt(8)
            h.paragraph_format.space_after = Pt(4)
        elif line_strip.startswith('##'):
            h = doc.add_heading(line_strip.replace('##', '').strip(), level=2)
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(6)
        elif line_strip.startswith('#'):
            h = doc.add_heading(line_strip.replace('#', '').strip(), level=1)
            h.paragraph_format.space_before = Pt(16)
            h.paragraph_format.space_after = Pt(8)
        # Bullet Lists
        elif line_strip.startswith('*') or line_strip.startswith('-'):
            clean_li = re.sub(r'^[\*\-]\s*', '', line_strip)
            doc.add_paragraph(clean_li, style='List Bullet')
        # Standard Paragraphs
        else:
            doc.add_paragraph(line_strip)
            
    doc.save(file_path)


MIRO_FALLBACK_BOARD_URL = "https://miro.com/app/board/uXjVPkMck5Q=/"


def _miro_headers(miro_token):
    return {
        "Authorization": f"Bearer {miro_token}",
        "Content-Type": "application/json",
        "Accept": "application/json"
    }


def _miro_create_board(headers, name, description):
    r = requests.post("https://api.miro.com/v2/boards", headers=headers,
                       json={"name": name, "description": description}, timeout=10)
    if r.status_code not in (200, 201):
        return None, None
    data = r.json()
    return data.get("id"), data.get("viewLink")


def _miro_create_text(headers, board_id, html, x, y, width, font_size="32", color="#e2e8f0", align="center"):
    requests.post(f"https://api.miro.com/v2/boards/{board_id}/texts", headers=headers, json={
        "data": {"content": html},
        "style": {"fontSize": font_size, "color": color, "textAlign": align},
        "position": {"x": x, "y": y},
        "geometry": {"width": width}
    }, timeout=8)


def _miro_create_frame(headers, board_id, title, fill_color, x, y, width, height):
    r = requests.post(f"https://api.miro.com/v2/boards/{board_id}/frames", headers=headers, json={
        "data": {"title": title, "format": "custom", "type": "freeform"},
        "style": {"fillColor": fill_color},
        "position": {"x": x, "y": y},
        "geometry": {"width": width, "height": height}
    }, timeout=8)
    return r.json().get("id") if r.status_code in (200, 201) else None


def _miro_create_shape(headers, board_id, shape, content, fill, border, text_color,
                        x, y, width, height, parent_id=None, fallback_x=None, fallback_y=None):
    payload = {
        "data": {"content": content, "shape": shape},
        "style": {"fillColor": fill, "borderColor": border, "color": text_color, "fontSize": "12"},
        "position": {"x": x, "y": y},
        "geometry": {"width": width, "height": height}
    }
    if parent_id:
        payload["parent"] = {"id": parent_id}
    r = requests.post(f"https://api.miro.com/v2/boards/{board_id}/shapes", headers=headers, json=payload, timeout=8)
    if r.status_code not in (200, 201) and parent_id and fallback_x is not None:
        # Nested placement can be rejected if it lands outside the frame's local
        # bounds - retry once as an absolute (unparented) canvas position instead.
        payload.pop("parent", None)
        payload["position"] = {"x": fallback_x, "y": fallback_y}
        r = requests.post(f"https://api.miro.com/v2/boards/{board_id}/shapes", headers=headers, json=payload, timeout=8)
    return r.json().get("id") if r.status_code in (200, 201) else None


def _miro_create_connector(headers, board_id, start_id, end_id, color, label=None):
    payload = {
        "startItem": {"id": start_id},
        "endItem": {"id": end_id},
        "shape": "curved",
        "style": {"strokeColor": color, "strokeWidth": "2", "startStrokeCap": "none", "endStrokeCap": "stealth"}
    }
    if label:
        payload["captions"] = [{"content": label, "position": "50%"}]
    requests.post(f"https://api.miro.com/v2/boards/{board_id}/connectors", headers=headers, json=payload, timeout=8)


LANE_ACCENTS = ["#10b981", "#3b82f6", "#a855f7", "#f59e0b", "#ec4899"]
LANE_FILLS = ["#0b2b22", "#0b1f33", "#1f1533", "#332708", "#330d20"]

COL_SPACING = 340
LEFT_MARGIN = 170
ROW_HEIGHT = 160
LANE_PAD = 40
LANE_GAP = 60


def _compute_flow_columns(steps):
    """Longest-path-from-root column assignment so branches/merges never overlap."""
    by_id = {s["id"]: s for s in steps}
    edges = []
    for s in steps:
        sid = s["id"]
        if s.get("type") == "decision":
            if s.get("yes_next") in by_id:
                edges.append((sid, s["yes_next"], "yes", s.get("yes_label") or "Yes"))
            if s.get("no_next") in by_id:
                edges.append((sid, s["no_next"], "no", s.get("no_label") or "No"))
        elif s.get("next") in by_id:
            edges.append((sid, s["next"], "default", None))

    indegree = {sid: 0 for sid in by_id}
    children = {sid: [] for sid in by_id}
    for f, t, _, _ in edges:
        children[f].append(t)
        indegree[t] += 1

    columns = {sid: 0 for sid in by_id if indegree[sid] == 0}
    frontier = list(columns.keys())
    guard = 0
    while frontier and guard < len(by_id) + 5:
        guard += 1
        nxt = []
        for sid in frontier:
            for c in children.get(sid, []):
                new_col = columns[sid] + 1
                if columns.get(c, -1) < new_col:
                    columns[c] = new_col
                    nxt.append(c)
        frontier = nxt
    for sid in by_id:
        columns.setdefault(sid, 0)
    return columns, edges


def create_miro_process_flow(miro_token, board_name, data):
    """
    Builds an enterprise-style swimlane process flow diagram: one lane frame per
    actor/system, shape-coded steps (pill start/end, rectangle process, diamond
    decision), colour-coded Yes/No branch connectors, a title, and a legend.
    """
    if not miro_token or "test_token" in miro_token or len(miro_token) < 20:
        return MIRO_FALLBACK_BOARD_URL
    steps = data.get("steps") or []
    if not steps:
        return MIRO_FALLBACK_BOARD_URL

    try:
        headers = _miro_headers(miro_token)
        board_id, board_url = _miro_create_board(headers, board_name, "Enterprise process flow diagram generated by Project War Room")
        if not board_id:
            return MIRO_FALLBACK_BOARD_URL

        lanes = list(data.get("lanes") or [])
        if not lanes:
            lanes = ["Process"]
        lane_index = {name: i for i, name in enumerate(lanes)}
        for s in steps:
            lane = s.get("lane") or lanes[0]
            s["lane"] = lane
            if lane not in lane_index:
                lane_index[lane] = len(lanes)
                lanes.append(lane)

        columns, edges = _compute_flow_columns(steps)

        # A (lane, column) cell can hold more than one node - e.g. two
        # different failure branches that both land back in the same lane one
        # hop after a decision. Without a sub-row, those nodes would share the
        # exact same coordinates and draw stacked exactly on top of each
        # other, which is what made every earlier diagram look like a single
        # flat row: only the last shape drawn per cell was ever visible.
        cell_occupancy = {}
        sub_row = {}
        for s in steps:
            cell = (s["lane"], columns[s["id"]])
            row = cell_occupancy.get(cell, 0)
            sub_row[s["id"]] = row
            cell_occupancy[cell] = row + 1

        lane_rows = {name: 1 for name in lanes}
        for (lane, _col), count in cell_occupancy.items():
            lane_rows[lane] = max(lane_rows[lane], count)
        lane_heights = {name: lane_rows[name] * ROW_HEIGHT + LANE_PAD for name in lanes}

        lane_y_offset = {}
        cursor = 0
        for name in lanes:
            lane_y_offset[name] = cursor
            cursor += lane_heights[name] + LANE_GAP
        total_lanes_height = cursor - LANE_GAP

        max_col = max(columns.values()) if columns else 0
        frame_width = (max_col + 1) * COL_SPACING + 200

        _miro_create_text(headers, board_id, f"<p><strong>{board_name}</strong></p>",
                           frame_width / 2, -220, frame_width, font_size="40")

        frame_ids = {}
        for name in lanes:
            fy = lane_y_offset[name] + lane_heights[name] / 2
            frame_ids[name] = _miro_create_frame(
                headers, board_id, f"\U0001F464 {name}", LANE_FILLS[lane_index[name] % len(LANE_FILLS)],
                frame_width / 2, fy, frame_width, lane_heights[name]
            )

        # BPMN-flavoured shapes: small circles for start/end events, a rounded
        # "task" box for process steps, a diamond for decisions/gateways.
        NODE_STYLES = {
            "start": ("circle", "#065f46", "#10b981", "#ffffff", 130, 130, True),
            "end": ("circle", "#065f46", "#10b981", "#ffffff", 130, 130, True),
            "end_success": ("circle", "#065f46", "#10b981", "#ffffff", 130, 130, True),
            "end_failure": ("circle", "#7f1d1d", "#ef4444", "#ffffff", 130, 130, True),
            "decision": ("rhombus", "#78350f", "#f59e0b", "#ffffff", 260, 160, False),
        }

        shape_ids = {}
        for s in steps:
            sid = s["id"]
            lane = s["lane"]
            col = columns.get(sid, 0)
            row = sub_row[sid]
            accent = LANE_ACCENTS[lane_index.get(lane, 0) % len(LANE_ACCENTS)]
            shape, fill, border, txt_color, w, h, terse = NODE_STYLES.get(
                s.get("type", "process"), ("round_rectangle", "#151f32", accent, "#e2e8f0", 240, 130, False)
            )
            title = s.get("title", f"Step {sid}")
            text = s.get("text", "")
            content = f"<p><strong>{title}</strong></p>" if terse else f"<p><strong>{title}</strong></p><p>{text}</p>"

            local_x = LEFT_MARGIN + col * COL_SPACING
            local_y = LANE_PAD + (row + 0.5) * ROW_HEIGHT
            abs_x = col * COL_SPACING + LEFT_MARGIN
            abs_y = lane_y_offset[lane] + local_y

            shape_ids[sid] = _miro_create_shape(
                headers, board_id, shape, content, fill, border, txt_color,
                local_x, local_y, w, h, parent_id=frame_ids.get(lane),
                fallback_x=abs_x, fallback_y=abs_y
            )

        for frm, to, kind, label in edges:
            if shape_ids.get(frm) and shape_ids.get(to):
                color = {"default": "#64748b", "yes": "#22c55e", "no": "#ef4444"}.get(kind, "#64748b")
                _miro_create_connector(headers, board_id, shape_ids[frm], shape_ids[to], color, label)

        legend_y = total_lanes_height + 60
        _miro_create_text(headers, board_id, "<p><strong>Legend</strong></p>", 60, legend_y, 200,
                           font_size="20", color="#94a3b8", align="left")
        legend_items = [
            ("circle", "#065f46", "#10b981", "Start / Success End"),
            ("round_rectangle", "#151f32", "#3b82f6", "Process Step"),
            ("rhombus", "#78350f", "#f59e0b", "Decision Point"),
            ("circle", "#7f1d1d", "#ef4444", "Failure / Rejected End"),
        ]
        lx = 0
        for shape, fill, border, label in legend_items:
            _miro_create_shape(headers, board_id, shape, "", fill, border, "#ffffff", lx, legend_y + 80, 90, 50)
            _miro_create_text(headers, board_id, f"<p>{label}</p>", lx, legend_y + 130, 220,
                               font_size="14", color="#94a3b8", align="left")
            lx += 260

        return board_url
    except Exception:
        return MIRO_FALLBACK_BOARD_URL


TIER_ACCENTS = ["#3b82f6", "#a855f7", "#0d9488", "#f59e0b", "#16a34a"]
TIER_FILLS = ["#0b1f33", "#1f1533", "#062b28", "#332708", "#0b2b1c"]

TIER_COMP_SPACING = 300
TIER_LEFT_MARGIN = 150
TIER_HEIGHT = 220
TIER_GAP = 90


def create_miro_architecture_diagram(miro_token, board_name, data):
    """
    Builds an enterprise-style layered system architecture diagram: one frame per
    tier (Presentation/Gateway/Service/Data, etc.), colour-coded per tier, with
    labelled connectors showing the protocol/data exchanged between components.
    """
    if not miro_token or "test_token" in miro_token or len(miro_token) < 20:
        return MIRO_FALLBACK_BOARD_URL
    tiers = data.get("tiers") or []
    if not tiers:
        return MIRO_FALLBACK_BOARD_URL

    try:
        headers = _miro_headers(miro_token)
        board_id, board_url = _miro_create_board(headers, board_name, "Enterprise system architecture diagram generated by Project War Room")
        if not board_id:
            return MIRO_FALLBACK_BOARD_URL

        max_components = max((len(t.get("components") or []) for t in tiers), default=1) or 1
        frame_width = max_components * TIER_COMP_SPACING + 200

        _miro_create_text(headers, board_id, f"<p><strong>{board_name}</strong></p>",
                           frame_width / 2, -220, frame_width, font_size="40")

        frame_ids = []
        shape_ids = []
        for t_idx, tier in enumerate(tiers):
            accent = TIER_ACCENTS[t_idx % len(TIER_ACCENTS)]
            fill = TIER_FILLS[t_idx % len(TIER_FILLS)]
            fy = t_idx * (TIER_HEIGHT + TIER_GAP) + TIER_HEIGHT / 2
            frame_id = _miro_create_frame(headers, board_id, f"\U0001F5C2 {tier.get('name', f'Tier {t_idx+1}')}",
                                           fill, frame_width / 2, fy, frame_width, TIER_HEIGHT)
            frame_ids.append(frame_id)

            components = tier.get("components") or []
            tier_shape_ids = []
            for c_idx, comp in enumerate(components):
                title = comp.get("title", f"Component {c_idx+1}")
                text = comp.get("text", "")
                content = f"<p><strong>⚙ {title}</strong></p><p>{text}</p>"
                local_x = TIER_LEFT_MARGIN + c_idx * TIER_COMP_SPACING
                local_y = TIER_HEIGHT / 2
                abs_x = c_idx * TIER_COMP_SPACING + TIER_LEFT_MARGIN
                abs_y = t_idx * (TIER_HEIGHT + TIER_GAP) + TIER_HEIGHT / 2
                sid = _miro_create_shape(
                    headers, board_id, "rectangle", content, "#151f32", accent, "#e2e8f0",
                    local_x, local_y, 260, 130, parent_id=frame_id,
                    fallback_x=abs_x, fallback_y=abs_y
                )
                tier_shape_ids.append(sid)
            shape_ids.append(tier_shape_ids)

        connections = data.get("connections") or []
        valid_connections = []
        for c in connections:
            ft, fc, tt, tc = c.get("from_tier"), c.get("from_component"), c.get("to_tier"), c.get("to_component")
            if (isinstance(ft, int) and isinstance(tt, int) and 0 <= ft < len(shape_ids) and 0 <= tt < len(shape_ids)
                    and isinstance(fc, int) and isinstance(tc, int)
                    and 0 <= fc < len(shape_ids[ft]) and 0 <= tc < len(shape_ids[tt])):
                valid_connections.append(c)

        if not valid_connections:
            # Fall back to a straightforward tier-to-tier cascade so the diagram
            # never renders as a disconnected pile of boxes.
            for t_idx in range(len(tiers) - 1):
                for c_idx in range(len(shape_ids[t_idx])):
                    target = min(c_idx, len(shape_ids[t_idx + 1]) - 1) if shape_ids[t_idx + 1] else None
                    if target is not None:
                        valid_connections.append({"from_tier": t_idx, "from_component": c_idx,
                                                    "to_tier": t_idx + 1, "to_component": target, "label": None})

        for c in valid_connections:
            start_id = shape_ids[c["from_tier"]][c["from_component"]]
            end_id = shape_ids[c["to_tier"]][c["to_component"]]
            if start_id and end_id:
                accent = TIER_ACCENTS[c["from_tier"] % len(TIER_ACCENTS)]
                _miro_create_connector(headers, board_id, start_id, end_id, accent, c.get("label"))

        legend_y = len(tiers) * (TIER_HEIGHT + TIER_GAP) + 60
        _miro_create_text(headers, board_id, "<p><strong>Legend</strong></p>", 60, legend_y, 200,
                           font_size="20", color="#94a3b8", align="left")
        lx = 0
        for t_idx, tier in enumerate(tiers):
            accent = TIER_ACCENTS[t_idx % len(TIER_ACCENTS)]
            _miro_create_shape(headers, board_id, "rectangle", "", "#151f32", accent, "#ffffff", lx, legend_y + 80, 90, 50)
            _miro_create_text(headers, board_id, f"<p>{tier.get('name', f'Tier {t_idx+1}')}</p>", lx, legend_y + 130, 220,
                               font_size="14", color="#94a3b8", align="left")
            lx += 260

        return board_url
    except Exception:
        return MIRO_FALLBACK_BOARD_URL


ARTIFACT_KEYWORDS = [
    ("User Story Map & Sprint Backlog", ["jira", "user stor", "sprint", "backlog", "epic", "subtask", "task", "work item", "work breakdown"]),
    ("Process Flow Diagram (Miro-designed)", ["process flow", "user journey", "flowchart", "flow diagram"]),
    ("System Architecture Specification (Miro-designed)", ["architecture", "system design", "tech stack", "infrastructure", "system diagram"]),
    ("Validation Test Scripts (UAT)", ["test script", "test case", "uat", "qa test", "validation test"]),
    ("Business Requirements Document (BRD)", ["brd", "business requirement"]),
    ("Functional Requirements Specification (FRS)", ["frs", "functional requirement"]),
    ("Use Case Specification", ["use case"]),
]
ALL_ARTIFACTS = [name for name, _ in ARTIFACT_KEYWORDS]

# The 3 contexts the app organizes work into: written documentation (synced to
# Confluence), Jira user stories/tasks (drafted then pushed to Jira), and Miro
# diagrams. Offered as quick-pick options when a request is too generic to
# scope on its own (e.g. "create documentation for the requirement").
CONTEXT_BUCKETS = {
    "documentation": ["Business Requirements Document (BRD)", "Functional Requirements Specification (FRS)", "Use Case Specification", "Validation Test Scripts (UAT)"],
    "jira": ["User Story Map & Sprint Backlog"],
    "miro": ["Process Flow Diagram (Miro-designed)", "System Architecture Specification (Miro-designed)"],
}
CONTEXT_BUCKETS["all"] = ALL_ARTIFACTS


def detect_target_artifacts(text):
    """
    Keyword-scoped artifact detection, so a targeted request ("rework the
    Jira tasks alone") only regenerates what it names instead of the whole
    document set. Deliberately not an LLM call - the chat path already has
    known hang risk from the Gemini SDK call having no timeout, and this
    doesn't need one: the artifact set is a small fixed list with distinctive
    names.

    Only scans the CURRENT message, never the accumulated chat history: an
    earlier confirmation like "we'll prepare the Business Requirements
    Document..." otherwise stays in history_str forever and would make every
    later, unrelated message in the same conversation spuriously re-match BRD.

    Returns [] (not a default) when nothing specific is named - the caller
    is expected to ask the user to pick a context rather than silently
    guessing the full document set.
    """
    lowered = text.lower()
    return [name for name, keywords in ARTIFACT_KEYWORDS if any(kw in lowered for kw in keywords)]


BABOK_INDEX_URL = "https://raw.githubusercontent.com/Prady089/BABOK_BA_Techniques_Handbook/main/BABOK%C2%AE%20Business%20Analysis%20Techniques%20%E2%80%93%20Master%20Index.md"
_babok_index_cache = {"content": None}


def get_babok_index():
    """
    Fetches and caches the BABOK v3 technique master index (~7KB table of
    all 50 named techniques, grouped by category) so agents can ground
    recommendations in real named techniques instead of generic advice.
    Cached in-memory for the life of the process - it's a static reference
    doc, not something that changes mid-session. A fetch failure (GitHub
    unreachable, etc.) just means callers proceed without it rather than
    blocking a chat response on external availability.
    """
    if _babok_index_cache["content"] is not None:
        return _babok_index_cache["content"]
    try:
        r = requests.get(BABOK_INDEX_URL, timeout=8)
        if r.status_code == 200:
            _babok_index_cache["content"] = r.text
            return r.text
    except Exception:
        pass
    return ""


def build_kb_context(knowledge_list):
    """
    Shared by /chat/ba, /chat/simulate_agents, and /chat/generate_artifacts:
    folds the project's manually-added knowledge items and the BABOK
    technique index into one context block for the LLM prompt.
    """
    context = ""
    if knowledge_list:
        context = "\nACTIVE KNOWLEDGE BASE REFERENCES:\n"
        for k in knowledge_list:
            context += f"- Title: {k.get('title')}\n  Context: {k.get('content')}\n"
    babok_index = get_babok_index()
    if babok_index:
        context += (
            "\nBABOK® v3 TECHNIQUE REFERENCE - the team's default methodology grounding. "
            "Where relevant, cite specific named techniques from this index (e.g. \"per BABOK's "
            f"Business Rules Analysis technique...\") rather than generic advice:\n{babok_index}\n"
        )
    return context


# ============================================================
# INTERACTIVE BA CHAT INTERVIEW ENDPOINT
# ============================================================

@app.post("/chat/ba")
async def chat_ba(request: Request):
    body = await request.json()
    user_message = body.get("message", "")
    history = body.get("history", [])
    knowledge_list = body.get("knowledge", [])
    project_id = body.get("project_id")

    kb_context = build_kb_context(knowledge_list)

    history_str = ""
    user_turns_count = 0
    for h in history:
        role = "User" if h["role"] == "user" else "BA Agent"
        history_str += f"{role}: {h['text']}\n"
        if h["role"] == "user":
            user_turns_count += 1

    if user_turns_count >= 1:
        # No agent-selection or artifact-checklist gate: the full specialist
        # team always runs by default, and the artifact set is inferred from
        # what was actually asked for (e.g. "rework the Jira tasks" only
        # regenerates the work item tree, not the entire document set) - see
        # detect_target_artifacts. This keeps the chat task-oriented instead
        # of forcing every request through the same multi-step wizard.
        target_artifacts = detect_target_artifacts(user_message)

        if not target_artifacts:
            # Too generic to scope on its own (e.g. "create documentation for
            # the requirement", or a fresh ask with no artifact named at all) -
            # ask which of the 3 contexts instead of silently guessing the
            # full document set. The frontend renders this as quick-pick
            # buttons; the user can also just type exactly what they want.
            response_text = (
                "Sure — what should the team prepare: Documentation (BRD, FRS, Use Cases & Test Scripts), "
                "Jira User Stories, Miro Diagrams, or all three? You can also just tell me exactly what you need."
            )
            if project_id:
                try:
                    db.add_chat_message(project_id, "user", user_message)
                    db.add_chat_message(project_id, "agent", response_text)
                except Exception:
                    pass
            return JSONResponse({
                "response": response_text,
                "requires_approval": False,
                "stage": "clarify",
                "context_buckets": CONTEXT_BUCKETS
            })

        scope_note = "the full requirements document set" if target_artifacts == ALL_ARTIFACTS else ", ".join(target_artifacts)
        system_prompt = f"""
You are acting as an expert Business Analyst (BA) in a War Room. The specialist team (Product Owner, UI/UX Designer, Systems Architect, Lead Developer, QA Engineer) will now automatically review this request and prepare: {scope_note}.
Write a short 1-2 sentence confirmation of what will be produced, in a natural conversational tone. Do not ask the user to select or approve anything - just confirm the scope and that the team is getting to work now.
{kb_context}
"""
        prompt = f"Chat History:\n{history_str}\n\nUser New Input: {user_message}\n\nConfirm scope and that the team is starting now:"
        response_text = llm(system_prompt, prompt)
        if project_id:
            try:
                db.add_chat_message(project_id, "user", user_message)
                db.add_chat_message(project_id, "agent", response_text)
            except Exception:
                pass
        return JSONResponse({
            "response": response_text,
            "requires_approval": False,
            "stage": "ready",
            "target_artifacts": target_artifacts
        })

    system_prompt = f"""
You are acting as an expert Business Analyst (BA).
The user wants to describe a requirement. Ask exactly 1 or 2 high-level multiple choice questions to establish the baseline requirement boundaries.
Keep it extremely simple. Do not ask detailed technical details or deep functional edge cases.
Always include "Other (please specify below)" as the final option.
{kb_context}
"""
    prompt = f"User requirement: {user_message}\n\nPlease ask 1-2 high-level baseline questions:"
    response_text = llm(system_prompt, prompt)
    if project_id:
        try:
            db.add_chat_message(project_id, "user", user_message)
            db.add_chat_message(project_id, "agent", response_text)
        except Exception:
            pass
    return JSONResponse({
        "response": response_text,
        "requires_approval": False,
        "stage": "interview"
    })


# ============================================================
# AGENT DELEGATION / DISCUSSION SIMULATION ENDPOINT
# ============================================================

@app.post("/chat/simulate_agents")
async def simulate_agents(request: Request):
    body = await request.json()
    requirements_history = body.get("history", [])
    agent_name = body.get("agent_name", "Architect")
    knowledge_list = body.get("knowledge", [])
    artifacts_list = body.get("artifacts", [])
    project_id = body.get("project_id")

    kb_context = build_kb_context(knowledge_list)

    req_summary = ""
    for h in requirements_history:
        role = "User" if h["role"] == "user" else "BA Agent"
        req_summary += f"{role}: {h['text']}\n"

    artifacts_str = ", ".join(artifacts_list) if artifacts_list else "selected documents"

    system_prompt = f"""
You are simulating a design room chat. You are the **{agent_name}**.
You are discussing with the lead **Business Analyst (BA)** to finalize the plan for the following requested artifacts: {artifacts_str}.

Provide 2-4 concrete, specific contribution points from your role's expertise that MUST be reflected in the final generated documents.
Each point must be a specific requirement, constraint, edge case, or decision grounded in the actual requirement discussed below — not generic boilerplate advice.
Format as short bullet points prefixed with "-". Keep the tone conversational in the intro line, but the bullet points must be concrete enough that a document writer could act on them directly.
If any selected document is a 'Process Flow Diagram' or 'System Architecture Specification', mention that we will use Miro to design the flow.
{kb_context}
"""

    prompt = f"Requirements history:\n{req_summary}\n\nPlease share your perspective as the {agent_name}:"
    response_text = llm(system_prompt, prompt)
    if project_id:
        try:
            db.add_agent_contribution(project_id, agent_name, response_text)
        except Exception:
            pass
    return JSONResponse({"transcript": response_text})


# ============================================================
# ARTIFACT GENERATION ENDPOINT
# ============================================================

@app.post("/chat/generate_artifacts")
async def generate_artifacts(request: Request):
    body = await request.json()
    requirements_history = body.get("history", [])
    knowledge_list = body.get("knowledge", [])
    artifacts_requested = body.get("artifacts", [])
    agent_contributions = body.get("agent_contributions", [])
    project_dir, safe_project_id = get_project_dir(body.get("project_id"))
    project_display_name = "Requirements"
    if body.get("project_id"):
        try:
            proj_basic = db.get_project_basic(body.get("project_id"))
            if proj_basic and proj_basic.get("name"):
                project_display_name = proj_basic["name"]
        except Exception:
            pass

    if not artifacts_requested:
        artifacts_requested = ["Business Requirements Document (BRD)", "System Architecture Spec", "Validation Test Scripts (UAT)"]

    kb_context = build_kb_context(knowledge_list)

    team_context = ""
    if agent_contributions:
        team_context = "\nTEAM INPUT FROM SPECIALIST AGENTS (these points MUST be reflected in the generated documents where relevant to that document's scope):\n"
        for c in agent_contributions:
            team_context += f"\n--- {c.get('agent', 'Team Member')} ---\n{c.get('text', '')}\n"

    req_summary = ""
    for h in requirements_history:
        role = "User" if h["role"] == "user" else "BA Agent"
        req_summary += f"{role}: {h['text']}\n"

    # Prepend knowledge base and specialist team context to requirements summary
    if kb_context:
        req_summary = kb_context + "\n" + req_summary
    if team_context:
        req_summary = team_context + "\n" + req_summary

    files_list = []
    work_item_tree = []
    generation_errors = []  # doc names skipped because llm() returned a failure (quota/API error) instead of content

    # 1. Generate Business Requirements Document (BRD)
    if "Business Requirements Document (BRD)" in artifacts_requested:
        brd_sys = "You are a Business Analyst. Generate a detailed Business Requirements Document (BRD) in Markdown format. Structure: 1. Executive Summary, 2. Scope, 3. Functional Requirements, 4. Non-Functional Requirements, 5. User Stories."
        brd_content = llm(brd_sys, req_summary)
        if _is_llm_error(brd_content):
            generation_errors.append("BRD.docx")
        else:
            # Save as Markdown
            with open(os.path.join(project_dir, "BRD.md"), "w", encoding="utf-8") as f:
                f.write(brd_content)
            # Save as DOCX
            docx_brd_path = os.path.join(project_dir, "BRD.docx")
            save_as_docx("Business Requirements Document (BRD)", brd_content, docx_brd_path)
            files_list.append({"name": "BRD.docx", "url": f"/workspace/{safe_project_id}/BRD.docx", "content": brd_content})

    # 2. Generate Functional Requirements Specification (FRS)
    if "Functional Requirements Specification (FRS)" in artifacts_requested:
        frs_sys = "You are a Functional Analyst. Generate a detailed Functional Requirements Specification (FRS) including system behaviors, error handling, and state transitions."
        frs_content = llm(frs_sys, req_summary)
        if _is_llm_error(frs_content):
            generation_errors.append("Functional_Requirements.docx")
        else:
            with open(os.path.join(project_dir, "Functional_Requirements.md"), "w", encoding="utf-8") as f:
                f.write(frs_content)
            docx_frs_path = os.path.join(project_dir, "Functional_Requirements.docx")
            save_as_docx("Functional Requirements Specification (FRS)", frs_content, docx_frs_path)
            files_list.append({"name": "Functional_Requirements.docx", "url": f"/workspace/{safe_project_id}/Functional_Requirements.docx", "content": frs_content})

    # 3. Generate Use Case Specification
    if "Use Case Specification" in artifacts_requested:
        uc_sys = "You are a Business Analyst. Generate a set of detailed Use Case Specifications including primary path, alternate paths, pre-conditions, and post-conditions."
        uc_content = llm(uc_sys, req_summary)
        if _is_llm_error(uc_content):
            generation_errors.append("Use_Cases.docx")
        else:
            with open(os.path.join(project_dir, "Use_Cases.md"), "w", encoding="utf-8") as f:
                f.write(uc_content)
            docx_uc_path = os.path.join(project_dir, "Use_Cases.docx")
            save_as_docx("Use Case Specification", uc_content, docx_uc_path)
            files_list.append({"name": "Use_Cases.docx", "url": f"/workspace/{safe_project_id}/Use_Cases.docx", "content": uc_content})

    # 4. Generate Process Flow Diagram (Miro-designed)
    if "Process Flow Diagram (Miro-designed)" in artifacts_requested:
        miro_token = os.getenv("MIRO_ACCESS_TOKEN", "eyJhbGciOiJIUzI1NiJ9.test_token")
        
        # 4a. Generate Word doc contents
        flow_sys = "You are a UI/UX Designer and Systems Analyst. Generate a detailed user journey and process flow specification including touchpoints, actions, and system responses."
        flow_content = llm(flow_sys, req_summary)
        if _is_llm_error(flow_content):
            generation_errors.append("Process_Flow_Diagram.docx")
        else:
            with open(os.path.join(project_dir, "Process_Flow_Diagram.md"), "w", encoding="utf-8") as f:
                f.write(flow_content)
            docx_flow_path = os.path.join(project_dir, "Process_Flow_Diagram.docx")
            save_as_docx("Process Flow Diagram Specification", flow_content, docx_flow_path)
            files_list.append({"name": "Process_Flow_Diagram.docx", "url": f"/workspace/{safe_project_id}/Process_Flow_Diagram.docx", "content": flow_content})

        # 4b. Parse a structured swimlane flow (actors, shape-typed steps, decision
        # branches) for an enterprise-grade Miro diagram - not a flat step list.
        flow_json_sys = """
You are a UI/UX Designer and Systems/Process Analyst. Design an enterprise-grade swimlane process flow diagram for the requirement below.

Output ONLY a single JSON object (no markdown, no backticks) with this exact shape:
{
  "lanes": ["Actor/System 1", "Actor/System 2"],
  "steps": [
    {"id": 1, "lane": "<one of the lanes above>", "type": "start", "title": "...", "text": "...", "next": 2},
    {"id": 2, "lane": "...", "type": "process", "title": "...", "text": "...", "next": 3},
    {"id": 3, "lane": "...", "type": "decision", "title": "...", "text": "...",
     "yes_label": "Valid", "yes_next": 4, "no_label": "Invalid", "no_next": 6},
    {"id": 4, "lane": "...", "type": "process", "title": "...", "text": "...", "next": 5},
    {"id": 5, "lane": "...", "type": "end_success", "title": "...", "text": "..."},
    {"id": 6, "lane": "...", "type": "end_failure", "title": "...", "text": "..."}
  ]
}
Rules:
- "lanes" must have 2-4 entries, one per distinct actor/system involved (e.g. "User", "API Gateway", "Auth Service").
- "type" must be one of: start, process, decision, end_success, end_failure.
- Every non-decision, non-end step must have a "next" id. Every "decision" step must have "yes_next"/"no_next" plus short "yes_label"/"no_label" (e.g. "Valid"/"Invalid", "Approved"/"Rejected").
- Include at least one decision point and at least one failure/rejection end path - real processes branch.
- "text" is a concise 1-2 sentence description of what happens at that step.
- Assign each step to the lane that actually performs it, so the diagram reads as a real cross-functional swimlane flow.
- Produce 6-12 steps total.
"""
        flow_raw = llm(flow_json_sys, req_summary)
        clean_flow_json = flow_raw.replace("```json", "").replace("```", "").strip()
        try:
            flow_data = json.loads(clean_flow_json)
            if not flow_data.get("steps"):
                raise ValueError("no steps")
        except Exception:
            flow_data = {
                "lanes": ["User", "Web/API Gateway", "Auth Service"],
                "steps": [
                    {"id": 1, "lane": "User", "type": "start", "title": "Arrives at Login Page", "text": "User navigates to the login page and enters credentials.", "next": 2},
                    {"id": 2, "lane": "Web/API Gateway", "type": "process", "title": "Submit Credentials", "text": "Gateway receives the request and forwards it to the Auth Service over HTTPS.", "next": 3},
                    {"id": 3, "lane": "Auth Service", "type": "decision", "title": "Validate Credentials", "text": "Checks username/password against the identity store.", "yes_label": "Valid", "yes_next": 4, "no_label": "Invalid", "no_next": 7},
                    {"id": 4, "lane": "Auth Service", "type": "decision", "title": "Device Recognized?", "text": "Checks device fingerprint and IP reputation.", "yes_label": "Trusted", "yes_next": 6, "no_label": "New Device", "no_next": 5},
                    {"id": 5, "lane": "Auth Service", "type": "process", "title": "Send MFA Challenge", "text": "Issues a one-time passcode via SMS/authenticator app.", "next": 6},
                    {"id": 6, "lane": "User", "type": "end_success", "title": "Access Granted", "text": "Secure session established; user redirected to dashboard."},
                    {"id": 7, "lane": "User", "type": "end_failure", "title": "Access Denied", "text": "Generic error shown; failed attempt logged for audit."}
                ]
            }

        # 4c. Create the swimlane flow diagram on a Miro board
        miro_board_url = create_miro_process_flow(miro_token, f"{project_display_name} - Process Flow Map", flow_data)
        files_list.append({"name": "Miro Process Flow Board (Interactive)", "url": miro_board_url, "external": True})

    # 5. Generate System Architecture Specification (Miro-designed)
    if "System Architecture Specification (Miro-designed)" in artifacts_requested or "System Architecture Spec" in artifacts_requested:
        miro_token = os.getenv("MIRO_ACCESS_TOKEN", "eyJhbGciOiJIUzI1NiJ9.test_token")
        
        arch_sys = """You are a Systems Architect. Generate a System Design and Architecture document.
Include a valid Mermaid diagram (inside standard markdown ```mermaid blocks) representing the authentication and backend service flows.
Make the diagram thorough and descriptive."""
        arch_content = llm(arch_sys, req_summary)
        if _is_llm_error(arch_content):
            generation_errors.append("System_Architecture.docx")
        else:
            with open(os.path.join(project_dir, "System_Architecture.md"), "w", encoding="utf-8") as f:
                f.write(arch_content)

            docx_arch_path = os.path.join(project_dir, "System_Architecture.docx")
            save_as_docx("System Architecture Specification", arch_content, docx_arch_path)
            files_list.append({"name": "System_Architecture.docx", "url": f"/workspace/{safe_project_id}/System_Architecture.docx", "content": arch_content})

        # 5b. Parse a structured layered architecture (tiers, components,
        # labelled inter-component connections) for an enterprise Miro diagram.
        arch_json_sys = """
You are a Systems/Enterprise Architect. Design an enterprise-grade layered system architecture diagram for the requirement below.

Output ONLY a single JSON object (no markdown, no backticks) with this exact shape:
{
  "tiers": [
    {"name": "Presentation Tier", "components": [{"title": "...", "text": "..."}]},
    {"name": "API / Gateway Tier", "components": [{"title": "...", "text": "..."}]},
    {"name": "Service Tier", "components": [{"title": "...", "text": "..."}]},
    {"name": "Data Tier", "components": [{"title": "...", "text": "..."}]}
  ],
  "connections": [
    {"from_tier": 0, "from_component": 0, "to_tier": 1, "to_component": 0, "label": "HTTPS/REST"}
  ]
}
Rules:
- Produce 3-5 tiers ordered top-to-bottom the way real traffic flows (e.g. Presentation -> Gateway -> Service(s) -> Data, with a Security/Cross-cutting tier if relevant).
- Each tier should have 1-3 concrete components (real service/component names appropriate to this requirement, not generic placeholders).
- "text" is a concise 1-2 sentence description of that component's responsibility.
- "connections" must reference valid 0-based tier/component indices from the "tiers" array above, and each "label" should name the protocol or data exchanged (e.g. "HTTPS/REST", "gRPC", "SQL", "Pub/Sub event").
- Include every meaningful connection, including lateral/cross-tier calls (e.g. Service Tier <-> Security Tier), not just a straight top-to-bottom chain.
"""
        arch_json_raw = llm(arch_json_sys, req_summary)
        clean_arch_json = arch_json_raw.replace("```json", "").replace("```", "").strip()
        try:
            arch_data = json.loads(clean_arch_json)
            if not arch_data.get("tiers"):
                raise ValueError("no tiers")
        except Exception:
            arch_data = {
                "tiers": [
                    {"name": "Presentation Tier", "components": [
                        {"title": "Web App (Browser)", "text": "Renders user dashboard and interfaces; initiates authentication requests."}
                    ]},
                    {"name": "API / Gateway Tier", "components": [
                        {"title": "API Gateway", "text": "Filters traffic, handles SSL termination, and routes requests to backend services."}
                    ]},
                    {"name": "Service Tier", "components": [
                        {"title": "Auth Microservice", "text": "Verifies credentials, checks password expiration, and manages sessions."},
                        {"title": "MFA Engine", "text": "Orchestrates TOTP codes, device OTP checks, and push validations."}
                    ]},
                    {"name": "Data Tier", "components": [
                        {"title": "Audit Logger & DB", "text": "Logs all events (IP, timestamp, status) to a high-availability database cluster."}
                    ]}
                ],
                "connections": [
                    {"from_tier": 0, "from_component": 0, "to_tier": 1, "to_component": 0, "label": "HTTPS/REST"},
                    {"from_tier": 1, "from_component": 0, "to_tier": 2, "to_component": 0, "label": "gRPC"},
                    {"from_tier": 2, "from_component": 0, "to_tier": 2, "to_component": 1, "label": "Internal call"},
                    {"from_tier": 2, "from_component": 0, "to_tier": 3, "to_component": 0, "label": "SQL/Write"},
                    {"from_tier": 2, "from_component": 1, "to_tier": 3, "to_component": 0, "label": "SQL/Write"}
                ]
            }

        miro_arch_url = create_miro_architecture_diagram(miro_token, f"{project_display_name} - System Architecture Map", arch_data)
        files_list.append({"name": "Miro System Architecture Board (Interactive)", "url": miro_arch_url, "external": True})

    # 6. Validation Test Scripts (UAT) no longer produces its own document -
    # "Validation Test Scripts (UAT)" as a requested artifact now only drives
    # step 8 below, which writes structured, Jira-traceable test cases
    # straight into the RTM (previously this also generated a free-text
    # Test_Scripts.docx from an independent LLM call, which routinely
    # disagreed with the RTM's own test cases since neither generation knew
    # about the other).

    # 7. Generate User Story Map & Sprint Backlog (Epic -> Story/Task -> Subtask)
    story_map_requested = "User Story Map & Sprint Backlog" in artifacts_requested or "User Story Map" in artifacts_requested
    if story_map_requested:
        stories_sys = """
You are a Product Owner. Generate a recommended Jira work breakdown for the final requirements, structured as a hierarchy.
This Jira site's hierarchy has exactly 3 levels: Epic at the top; Story and Task are PEERS directly under an Epic (never nest a Task under a Story); Subtask is the only level below, nested under a Story or a Task.

Format the output EXACTLY as a JSON array of objects, with no extra markdown styling or backticks. Each object has "issuetype", "summary", "description", "priority", and "children" (omit or use [] when there are none):
[
  {"issuetype": "Epic", "summary": "Epic summary", "description": "...", "priority": "High", "children": [
    {"issuetype": "Story", "summary": "Story summary", "description": "...", "priority": "High", "children": [
      {"issuetype": "Subtask", "summary": "Subtask summary", "description": "...", "priority": "Low"}
    ]},
    {"issuetype": "Task", "summary": "Task summary", "description": "...", "priority": "Medium", "children": []}
  ]}
]
A top-level entry may also be a standalone "Story" or "Task" with no Epic if it doesn't belong under one. Produce 1-3 Epics, each with a few Stories/Tasks, and Subtasks only where they add real value.
"""
        stories_content_raw = llm(stories_sys, req_summary)
        clean_json = stories_content_raw.replace("```json", "").replace("```", "").strip()
        try:
            work_item_tree = _sanitize_work_item_tree(json.loads(clean_json))
        except Exception:
            pass

        # Write a user stories docx file and add to files_list
        stories_docx_sys = "You are a Product Owner. Generate a detailed User Story Map & Sprint Backlog with User Stories, Tasks, and Acceptance Criteria in Markdown."
        stories_docx_content = llm(stories_docx_sys, req_summary)
        if _is_llm_error(stories_docx_content):
            generation_errors.append("User_Stories.docx")
        else:
            with open(os.path.join(project_dir, "User_Stories.md"), "w", encoding="utf-8") as f:
                f.write(stories_docx_content)
            docx_stories_path = os.path.join(project_dir, "User_Stories.docx")
            save_as_docx("User Story Map & Sprint Backlog", stories_docx_content, docx_stories_path)
            files_list.append({"name": "User_Stories.docx", "url": f"/workspace/{safe_project_id}/User_Stories.docx", "content": stories_docx_content})

        # Safety net for this artifact only - if the LLM's JSON came back
        # unparseable, still leave a usable starting draft rather than an
        # empty tree. Deliberately scoped inside `story_map_requested`: a
        # request that never asked for the story map must never touch the
        # project's Jira work items at all (see below).
        if not work_item_tree:
            work_item_tree = [
                {"issuetype": "Epic", "summary": "Secure Customer Authentication", "description": "Deliver MFA-based login for customers.", "priority": "High", "children": [
                    {"issuetype": "Story", "summary": "Implement login MFA component", "description": "As a customer, I want to authenticate securely via SMS/Email OTP.", "priority": "High", "children": []},
                    {"issuetype": "Task", "summary": "Configure system transaction audit logging", "description": "Log all transaction attempts to database.", "priority": "Medium", "children": []},
                ]},
            ]

    raw_project_id = body.get("project_id")
    if raw_project_id:
        try:
            # Merge with whatever the project already has instead of a full
            # replace - a scoped request (e.g. "just rework the process flow
            # diagram") only regenerates what it names via `files_list`, and
            # must not delete every other previously generated artifact.
            existing = db.get_project(raw_project_id)
            existing_artifacts = (existing or {}).get("artifacts", [])
            merged_by_name = {
                a["name"]: {"name": a["name"], "url": a["url"], "content": a.get("content"), "external": bool(a.get("is_external"))}
                for a in existing_artifacts
            }
            # No artifact generated by this endpoint is ever a test script
            # any more, but preserve the flag on anything already marked as
            # one (e.g. a legacy Test_Scripts.docx from before this change).
            existing_test_script_names = {a["name"] for a in existing_artifacts if a.get("is_test_script")}
            for f in files_list:
                merged_by_name[f["name"]] = f
            db.replace_artifacts(raw_project_id, list(merged_by_name.values()), existing_test_script_names)

            # Likewise, only touch work items when the story map was actually
            # part of this request - otherwise a scoped, unrelated request
            # would wipe the project's real draft Jira tasks and replace them
            # with the safety-net placeholder above.
            if story_map_requested:
                db.replace_unsynced_work_items(raw_project_id, work_item_tree)
            # Re-fetch from the DB so the response (and the RTM test-case step
            # below) reflect the project's true current tree either way - real
            # ids/codes for the frontend, and the actual persisted state for a
            # request that didn't touch work items at all.
            work_item_tree = db.get_work_item_tree(raw_project_id)
        except Exception:
            pass  # generation already succeeded and files are on disk; don't fail the response over a persistence hiccup

        # 8. Generate structured, Jira-traceable Test Cases for the RTM - each
        # one explicitly linked to a Story/Task code from the tree just persisted
        # above, so the Trace Matrix can show real coverage instead of a guess.
        if work_item_tree and ("Validation Test Scripts (UAT)" in artifacts_requested or "Validation Test Scripts (manual & automated cases)" in artifacts_requested):
            linkable_items = []

            def _collect_linkable(nodes):
                for n in nodes:
                    if n.get("issuetype") in ("Story", "Task") and n.get("code"):
                        linkable_items.append({"code": n["code"], "summary": n.get("summary")})
                    _collect_linkable(n.get("children") or [])

            _collect_linkable(work_item_tree)

            if linkable_items:
                tc_sys = f"""
You are a QA Engineer. Write concrete test cases that validate the following user stories/tasks. Write 1-2 test cases for EACH one listed - every code must be covered by at least one test case.

Stories/Tasks to cover:
{json.dumps(linkable_items, indent=2)}

Output ONLY a JSON array of objects (no markdown, no backticks):
[
  {{"linked_code": "STY001", "title": "Concise test case title", "steps": "1. ...\\n2. ...\\n3. ...", "expected_result": "...", "priority": "High"}}
]
"linked_code" MUST be exactly one of the codes listed above (case-sensitive, no other text).
"""
                tc_raw = llm(tc_sys, req_summary)
                clean_tc_json = tc_raw.replace("```json", "").replace("```", "").strip()
                try:
                    test_cases_data = json.loads(clean_tc_json)
                    db.replace_test_cases(raw_project_id, test_cases_data)
                except Exception:
                    pass

    return JSONResponse({
        "success": True,
        "files": files_list,
        "work_items": work_item_tree,
        "generation_errors": generation_errors,
    })


# ============================================================
# ARTIFACTS TAB "ASK AI" (generic doc create/update - not work items)
# ============================================================

def _merge_and_save_artifact(project_id, project_dir, safe_project_id, docx_name, title, content, existing_artifacts, is_test_script=False):
    """Writes one doc artifact (.md + .docx) to disk and full-replaces it into
    the project's artifact list, preserving every other existing artifact
    (Miro board links, other docs, test scripts) untouched. `existing_artifacts`
    must be freshly read from the DB (its rows carry the real is_test_script
    flag) - a list built from a previous call's return value does NOT, since
    those returned dicts don't carry that field, and would silently wipe the
    flag on every other doc if reused as the next call's base."""
    md_name = os.path.splitext(docx_name)[0] + ".md"
    with open(os.path.join(project_dir, md_name), "w", encoding="utf-8") as f:
        f.write(content)
    save_as_docx(title, content, os.path.join(project_dir, docx_name))

    merged_by_name = {
        a["name"]: {"name": a["name"], "url": a["url"], "content": a.get("content"), "external": bool(a.get("is_external"))}
        for a in existing_artifacts
    }
    test_script_names = {a["name"] for a in existing_artifacts if a.get("is_test_script")}
    if is_test_script:
        test_script_names.add(docx_name)
    else:
        test_script_names.discard(docx_name)
    merged_by_name[docx_name] = {"name": docx_name, "url": f"/workspace/{safe_project_id}/{docx_name}", "content": content}

    db.replace_artifacts(project_id, list(merged_by_name.values()), test_script_names)
    return list(merged_by_name.values())


@app.post("/artifacts/quick_update")
async def artifacts_quick_update(request: Request):
    """
    Artifacts tab's "Ask AI" box - a generic documentation assistant, scoped
    to documents only (never Jira work items, never Miro boards). Decides
    whether the request updates one of the project's EXISTING documents or
    writes a brand new one, then (re)writes the full document via the LLM.
    """
    body = await request.json()
    project_id = body.get("project_id")
    message = (body.get("message") or "").strip()
    if not project_id or not message:
        return JSONResponse({"error": "project_id and message are required"}, status_code=400)

    project = db.get_project(project_id)
    if not project:
        return JSONResponse({"error": "Project not found"}, status_code=404)

    existing_artifacts = project.get("artifacts", [])
    existing_docs = [a for a in existing_artifacts if a.get("content") and not a.get("is_external")]

    decide_sys = f"""
You are a documentation assistant for a Business Analysis workspace. The user will ask you to CREATE a new document or UPDATE an existing one - never Jira work items, never diagrams/boards.

Existing documents in this project:
{json.dumps([{"name": a["name"]} for a in existing_docs], indent=2) if existing_docs else "(none yet)"}

Output ONLY a JSON object (no markdown, no backticks):
{{"action": "update", "target_name": "<exact name copied from the list above>", "title": "<short display title>"}}
or
{{"action": "create", "title": "<short display title for the new document>"}}
Rules:
- Use "update" only when the request clearly targets a document already in the list (by name or obvious topic match). "target_name" MUST be copied exactly from the list, unchanged.
- Otherwise use "create".
"""
    decide_raw = llm(decide_sys, message)
    clean = decide_raw.replace("```json", "").replace("```", "").strip()
    try:
        decision = json.loads(clean)
        action = decision.get("action")
        title = (decision.get("title") or "New Document").strip()
        if action not in ("update", "create"):
            raise ValueError("bad action")
    except Exception:
        return JSONResponse({"error": "Could not understand that request - try rephrasing, e.g. \"Update the BRD to include...\" or \"Create a Data Migration Plan document\"."}, status_code=422)

    kb_context = build_kb_context(project.get("knowledge_items", []))
    project_dir, safe_project_id = get_project_dir(project_id)

    if action == "update":
        target = next((a for a in existing_docs if a["name"] == decision.get("target_name")), None)
        if not target:
            return JSONResponse({"error": "Couldn't match that to an existing document - try naming it explicitly, e.g. \"Update the BRD to...\"."}, status_code=422)
        write_sys = "You are revising an existing project document. Rewrite the COMPLETE document in Markdown, applying the requested change and keeping everything else that's still accurate. Output only the document itself, no commentary."
        user_prompt = f"{kb_context}\n\nCURRENT DOCUMENT ({target['name']}):\n{target.get('content') or ''}\n\nREQUESTED CHANGE:\n{message}"
        content = llm(write_sys, user_prompt)
        docx_name = target["name"]
        title = os.path.splitext(docx_name)[0].replace("_", " ")
        is_test_script = bool(target.get("is_test_script"))
    else:
        write_sys = "You are a Business Analyst writing project documentation. Produce a complete, well-structured document in Markdown for the request below."
        user_prompt = f"{kb_context}\n\nDOCUMENT REQUEST:\n{message}\n\nDocument title: {title}"
        content = llm(write_sys, user_prompt)
        safe_title = re.sub(r'[^A-Za-z0-9_\-]', '_', title)[:60] or "Document"
        existing_names_now = {a["name"] for a in existing_artifacts}
        docx_name = f"{safe_title}.docx"
        suffix = 1
        while docx_name in existing_names_now:
            suffix += 1
            docx_name = f"{safe_title}_{suffix}.docx"
        is_test_script = False

    if _is_llm_error(content):
        return JSONResponse({"error": f"Generation failed, nothing was changed: {content}"}, status_code=502)

    files = _merge_and_save_artifact(project_id, project_dir, safe_project_id, docx_name, title, content, existing_artifacts, is_test_script=is_test_script)
    return JSONResponse({"success": True, "action": action, "name": docx_name, "files": files})


# ============================================================
# UPDATE ARTIFACTS FROM LIVE JIRA
# ============================================================

def _adf_to_text(node):
    """Pulls plain text out of a Jira Cloud ADF (Atlassian Document Format)
    description field, ignoring formatting - good enough as LLM grounding
    context, not a faithful re-render."""
    parts = []
    def walk(n):
        if isinstance(n, dict):
            if n.get("type") == "text":
                parts.append(n.get("text", ""))
            for c in n.get("content") or []:
                walk(c)
        elif isinstance(n, list):
            for c in n:
                walk(c)
    walk(node)
    return " ".join(p for p in parts if p)


@app.post("/artifacts/update_from_jira")
async def update_artifacts_from_jira(request: Request):
    """
    Re-generates each document artifact that already exists (BRD, FRS, Use
    Cases, Test Scripts) so it reflects the CURRENT live Jira backlog. Closes
    the gap where a story added directly via Jira Sync's "Ask AI" (or in
    Jira itself) for a new requirement never flows back into the
    requirements docs, which are otherwise only written once from the War
    Room chat history. Only updates documents that already exist - it
    doesn't author brand-new ones (use the War Room chat or the Artifacts
    "Ask AI" box for that).
    """
    if not jira_configured():
        return JSONResponse({"error": "Jira not configured"}, status_code=503)
    body = await request.json()
    project_id = body.get("project_id")
    if not project_id:
        return JSONResponse({"error": "project_id is required"}, status_code=400)

    project = db.get_project(project_id)
    if not project:
        return JSONResponse({"error": "Project not found"}, status_code=404)

    jira_key, _ = resolve_project_keys(project_id)
    headers, auth = jira_headers()
    try:
        jql = f"project = {jira_key} ORDER BY created ASC"
        params = {"jql": jql, "maxResults": 200, "fields": "summary,description,issuetype,status,priority"}
        kwargs = {"headers": headers, "params": params, "verify": JIRA_SSL_VERIFY, "timeout": 15}
        if auth:
            kwargs["auth"] = auth
        r = requests.get(f"{JIRA_URL}/rest/api/3/search/jql", **kwargs)
        if r.status_code != 200:
            return JSONResponse({"error": f"Failed to read Jira: {r.text}"}, status_code=502)
        issues = r.json().get("issues", [])
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=502)

    if not issues:
        return JSONResponse({"error": "No issues found in Jira for this project yet"}, status_code=422)

    backlog_lines = []
    for issue in issues:
        f = issue.get("fields", {})
        desc_text = _adf_to_text(f.get("description")).strip()
        line = f"- [{issue['key']}] ({(f.get('issuetype') or {}).get('name', '')}) {f.get('summary', '')}"
        if desc_text:
            line += f": {desc_text}"
        backlog_lines.append(line)

    kb_context = build_kb_context(project.get("knowledge_items", []))
    req_summary = (
        f"{kb_context}\n\nCURRENT LIVE JIRA BACKLOG for project {jira_key} - this is the ground truth. "
        f"The document below MUST reflect every item listed here, including anything not previously covered:\n"
        + "\n".join(backlog_lines)
    )

    project_dir, safe_project_id = get_project_dir(project_id)
    existing_names = {a["name"] for a in project.get("artifacts", [])}

    doc_specs = [
        ("BRD.docx", "Business Requirements Document (BRD)", "You are a Business Analyst. Generate a detailed Business Requirements Document (BRD) in Markdown format. Structure: 1. Executive Summary, 2. Scope, 3. Functional Requirements, 4. Non-Functional Requirements, 5. User Stories."),
        ("Functional_Requirements.docx", "Functional Requirements Specification (FRS)", "You are a Functional Analyst. Generate a detailed Functional Requirements Specification (FRS) including system behaviors, error handling, and state transitions."),
        ("Use_Cases.docx", "Use Case Specification", "You are a Business Analyst. Generate a set of detailed Use Case Specifications including primary path, alternate paths, pre-conditions, and post-conditions."),
    ]

    updated = []
    failed = []
    files = None
    for docx_name, title, sys_prompt in doc_specs:
        if docx_name not in existing_names:
            continue  # only refresh docs the project already has
        content = llm(sys_prompt, req_summary)
        if _is_llm_error(content):
            # Never let a failed generation (quota exceeded, API error, etc.)
            # overwrite the existing document with the raw error text -
            # leave that doc untouched and report the failure instead.
            failed.append(docx_name)
            continue
        # Re-read from the DB each iteration so is_test_script/content on every
        # OTHER doc reflects the true current state, including what the
        # previous iteration of this same loop just wrote.
        current_artifacts = db.get_project(project_id)["artifacts"]
        is_test_script = docx_name in {a["name"] for a in current_artifacts if a.get("is_test_script")}
        files = _merge_and_save_artifact(project_id, project_dir, safe_project_id, docx_name, title, content, current_artifacts, is_test_script=is_test_script)
        updated.append(docx_name)

    if not updated and not failed:
        return JSONResponse({"error": "No existing documents to update yet - generate them first from the War Room chat."}, status_code=422)
    if not updated and failed:
        return JSONResponse({"error": f"Generation failed for every matching document ({', '.join(failed)}); nothing was changed."}, status_code=502)

    return JSONResponse({"success": True, "updated": updated, "failed": failed, "issue_count": len(issues), "files": files})


# ============================================================
# QUICK-ADD WORK ITEMS (AI natural-language + manual "+ Create")
# ============================================================

@app.post("/work_items/quick_create")
async def quick_create_work_items(request: Request):
    """
    Scoped, single-purpose assistant for the Draft pane's "Ask AI" box - NOT
    the general BA/agent chat. Takes a short natural-language request
    ("create a story for X under EP002, and 3 tasks under STY002 for...")
    and turns it into new draft work items attached under EXISTING items by
    code, additively - it never touches or regenerates anything else.
    """
    body = await request.json()
    project_id = body.get("project_id")
    message = (body.get("message") or "").strip()
    if not project_id or not message:
        return JSONResponse({"error": "project_id and message are required"}, status_code=400)

    try:
        existing = db.get_all_work_items(project_id)
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)

    existing_by_code = {i["code"]: i for i in existing if i.get("code")}
    reference_list = [
        {"code": i["code"], "issuetype": i["issuetype"], "summary": i["summary"]}
        for i in existing if i.get("code")
    ]

    sys_prompt = f"""
You are a Jira work-item assistant. The user will describe new Stories, Tasks, or Subtasks to add to an EXISTING breakdown - you are NOT writing requirements documents, just turning their request into structured new items.

Here is the current breakdown (code, type, summary) you can attach new items under:
{json.dumps(reference_list, indent=2)}

Output ONLY a JSON array (no markdown, no backticks) of insertions:
[
  {{"parent_code": "EP002", "node": {{"issuetype": "Story", "summary": "...", "description": "...", "priority": "Medium", "children": []}}}},
  {{"parent_code": "STY002", "node": {{"issuetype": "Task", "summary": "...", "description": "...", "priority": "Medium", "children": []}}}}
]
Rules:
- "parent_code" MUST be an exact code from the list above (e.g. "EP002", "STY002") when the user references an existing item ("under EPIC 002", "under story 002", etc - resolve loose phrasing to the exact matching code). Use null only for a genuinely new top-level Epic.
- Respect the real hierarchy: Epic has no parent; Story and Task attach to an Epic (or nothing); Subtask attaches only to a Story or a Task.
- If the user's request implies a NEW parent that doesn't exist yet (e.g. "create a story and 3 tasks under it"), nest the tasks inside that story's own "children" array in the SAME insertion instead of inventing a parent_code that doesn't exist.
- If the user asks for several of the same thing (e.g. "3 tasks"), emit that many separate Task nodes.
- Keep summaries concise and specific to what the user actually described.
"""
    raw = llm(sys_prompt, message)
    clean = raw.replace("```json", "").replace("```", "").strip()
    try:
        insertions_raw = json.loads(clean)
        if not isinstance(insertions_raw, list):
            raise ValueError("not a list")
    except Exception:
        return JSONResponse({"error": "Could not understand that request - try rephrasing, e.g. \"Create a story for X under EP002\"."}, status_code=422)

    valid_insertions = []
    skipped = []
    for ins in insertions_raw:
        if not isinstance(ins, dict) or not isinstance(ins.get("node"), dict):
            continue
        parent_code = ins.get("parent_code")
        parent_type = None
        if parent_code:
            parent = existing_by_code.get(parent_code)
            if not parent:
                skipped.append(f"Unknown parent code '{parent_code}'")
                continue
            parent_type = parent["issuetype"]
        cleaned_nodes = _sanitize_work_item_tree([ins["node"]], parent_type)
        if not cleaned_nodes:
            skipped.append(f"\"{ins['node'].get('summary', 'item')}\" doesn't fit under {parent_code or 'top level'}")
            continue
        valid_insertions.append({"parent_code": parent_code, "node": cleaned_nodes[0]})

    if not valid_insertions:
        return JSONResponse({
            "error": "Nothing valid to add" + (": " + "; ".join(skipped) if skipped else ""),
            "skipped": skipped
        }, status_code=422)

    try:
        created = db.add_work_items(project_id, valid_insertions)
        work_item_tree = db.get_work_item_tree(project_id)
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)

    return JSONResponse({"success": True, "created": created, "skipped": skipped, "work_items": work_item_tree})


@app.post("/work_items/manual_create")
async def manual_create_work_item(request: Request):
    """The Draft pane's "+ Create" form - a single item, no AI, structured input only."""
    body = await request.json()
    project_id = body.get("project_id")
    issuetype = body.get("issuetype")
    summary = (body.get("summary") or "").strip()
    parent_code = body.get("parent_code") or None
    priority = body.get("priority") or "Medium"
    description = body.get("description") or ""

    if not project_id or not summary:
        return JSONResponse({"error": "project_id and summary are required"}, status_code=400)
    if issuetype not in ("Epic", "Story", "Task", "Subtask"):
        return JSONResponse({"error": "Invalid issue type"}, status_code=400)

    parent_type = None
    if parent_code:
        parent = db.get_work_item_by_code(project_id, parent_code)
        if not parent:
            return JSONResponse({"error": f"Unknown parent code '{parent_code}'"}, status_code=400)
        parent_type = parent["issuetype"]

    node = {"issuetype": issuetype, "summary": summary, "description": description, "priority": priority, "children": []}
    cleaned = _sanitize_work_item_tree([node], parent_type)
    if not cleaned:
        return JSONResponse({"error": f"A {issuetype} can't go under a {parent_type or 'top level'} item"}, status_code=400)

    try:
        created = db.add_work_items(project_id, [{"parent_code": parent_code, "node": cleaned[0]}])
        work_item_tree = db.get_work_item_tree(project_id)
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)

    return JSONResponse({"success": True, "created": created, "work_items": work_item_tree})


# ============================================================
# TEST CASE GENERATION (per work item, for the RTM)
# ============================================================

@app.post("/work_items/{item_id}/generate_test_cases")
async def generate_test_cases_for_item(item_id: int):
    """
    Jira Sync panel's per-item "Generate Test Cases" action - writes 1-2
    concrete test cases for ONE Story/Task via the LLM and adds them to the
    local test_cases table, additively (never touches any other item's test
    cases). They immediately show up as covered rows in the RTM, in Draft
    status until pushed to Jira.
    """
    items = db.get_work_items_by_ids([item_id])
    if not items:
        return JSONResponse({"error": "Work item not found"}, status_code=404)
    item = items[0]
    if item["issuetype"] not in ("Story", "Task"):
        return JSONResponse({"error": "Test cases can only be generated for a Story or a Task"}, status_code=400)

    tc_sys = """
You are a QA Engineer. Write 1-2 concrete test cases that validate the Story/Task below.

Output ONLY a JSON array of objects (no markdown, no backticks):
[
  {"title": "Concise test case title", "steps": "1. ...\\n2. ...\\n3. ...", "expected_result": "...", "priority": "High"}
]
"""
    prompt = f"Summary: {item.get('summary')}\nDescription: {item.get('description') or '(none)'}"
    raw = llm(tc_sys, prompt)
    clean = raw.replace("```json", "").replace("```", "").strip()
    try:
        parsed = json.loads(clean)
        if not isinstance(parsed, list) or not parsed:
            raise ValueError("empty")
    except Exception:
        return JSONResponse({"error": "Could not generate test cases for this item - try again."}, status_code=422)

    try:
        created = db.add_test_cases(item["project_id"], item_id, parsed)
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)

    return JSONResponse({"success": True, "created": created})


# ============================================================
# PUSH TEST CASES TO JIRA (as linked Subtasks under their Story/Task)
# ============================================================

@app.post("/jira/push_test_cases")
async def jira_push_test_cases(request: Request):
    """
    Pushes selected test_cases rows to Jira. This site's hierarchy only has
    Epic/Story/Task/Subtask (no dedicated Test issue type - that requires an
    app like Xray or Zephyr), so a test case is created as a Subtask under
    its linked Story/Task, prefixed and labelled so it reads as a test case
    rather than a dev subtask. The parent Story/Task must already be synced.
    """
    if not jira_configured():
        return JSONResponse({"error": "Jira not configured"}, status_code=503)
    try:
        body = await request.json()
        project_id = body.get("project_id")
        test_case_ids = body.get("test_case_ids", [])
        if not project_id or not test_case_ids:
            return JSONResponse({"error": "project_id and test_case_ids are required"}, status_code=400)

        jira_key, _ = resolve_project_keys(project_id)
        headers, auth = jira_headers()

        test_cases = db.get_test_cases_by_ids(test_case_ids)
        work_items_by_id = {w["id"]: w for w in db.get_all_work_items(project_id)}

        results = []
        synced_updates = []
        for tc in test_cases:
            if tc.get("jira_key"):
                results.append({"id": tc["id"], "success": True, "jira_key": tc["jira_key"]})
                continue
            parent = work_items_by_id.get(tc.get("work_item_id"))
            parent_key = parent.get("jira_key") if parent else None
            if not parent_key:
                results.append({"id": tc["id"], "success": False, "error": "Parent Story/Task isn't pushed to Jira yet"})
                continue

            steps_text = (tc.get("steps") or "").strip()
            expected_text = (tc.get("expected_result") or "").strip()
            desc_lines = []
            if steps_text:
                desc_lines.append("Steps:\n" + steps_text)
            if expected_text:
                desc_lines.append("Expected Result:\n" + expected_text)

            payload = {
                "fields": {
                    "project": {"key": jira_key},
                    "summary": f"[TEST] {tc.get('title', 'Untitled Test Case')}",
                    "issuetype": {"name": "Subtask"},
                    "parent": {"key": parent_key},
                    "labels": ["test-case"],
                }
            }
            if desc_lines:
                payload["fields"]["description"] = {
                    "type": "doc",
                    "version": 1,
                    "content": [{"type": "paragraph", "content": [{"type": "text", "text": "\n\n".join(desc_lines)}]}]
                }
            if tc.get("priority"):
                payload["fields"]["priority"] = {"name": tc["priority"]}

            kwargs = {"headers": headers, "json": payload, "verify": JIRA_SSL_VERIFY, "timeout": 15}
            if auth:
                kwargs["auth"] = auth
            r = requests.post(f"{JIRA_URL}/rest/api/3/issue", **kwargs)
            if r.status_code in (200, 201):
                key = r.json().get("key")
                synced_updates.append({"id": tc["id"], "jira_key": key})
                results.append({"id": tc["id"], "success": True, "jira_key": key})
            else:
                results.append({"id": tc["id"], "success": False, "error": r.text})

        if synced_updates:
            try:
                db.mark_test_cases_synced(project_id, synced_updates)
            except Exception:
                pass

        return JSONResponse({"success": True, "results": results})
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)


# ============================================================
# PUSH WORK ITEMS TO JIRA (hierarchy-aware)
# ============================================================

@app.post("/jira/push_work_items")
async def jira_push_work_items(request: Request):
    """
    Pushes selected work_items rows to Jira, respecting the real 3-level
    hierarchy (Epic -> Story/Task -> Subtask) via Jira's `parent` field.
    `parent_overrides` (local work_item id -> jira key or null) is supplied
    by the frontend when the user picked a parent for an item whose real
    DB parent isn't selected/synced in this batch - if an item's parent
    still can't be resolved after that, it's reported as a failure rather
    than silently created as an orphan or sent to Jira and rejected there.
    """
    if not jira_configured():
        return JSONResponse({"error": "Jira not configured"}, status_code=503)
    try:
        body = await request.json()
        project_id = body.get("project_id")
        item_ids = body.get("item_ids", [])
        parent_overrides = {int(k): v for k, v in (body.get("parent_overrides") or {}).items()}

        if not project_id or not item_ids:
            return JSONResponse({"error": "project_id and item_ids are required"}, status_code=400)

        jira_key, _ = resolve_project_keys(project_id)
        headers, auth = jira_headers()

        all_items = {item["id"]: item for item in db.get_all_work_items(project_id)}
        pending = [all_items[i] for i in item_ids if i in all_items]

        resolved_jira_key = {item["id"]: item["jira_key"] for item in all_items.values() if item.get("jira_key")}
        results = []
        synced_updates = []

        remaining = pending
        for _ in range(4):  # 3 hierarchy levels + one safety margin
            if not remaining:
                break
            still_waiting = []
            for item in remaining:
                parent_id = item.get("parent_id")
                parent_key = None
                if parent_id is not None:
                    if parent_id in resolved_jira_key:
                        parent_key = resolved_jira_key[parent_id]
                    elif item["id"] in parent_overrides:
                        parent_key = parent_overrides[item["id"]]
                    else:
                        still_waiting.append(item)
                        continue
                elif item["id"] in parent_overrides and parent_overrides[item["id"]]:
                    parent_key = parent_overrides[item["id"]]

                if item["issuetype"] == "Subtask" and not parent_key:
                    results.append({"id": item["id"], "success": False, "error": "Subtask requires a parent"})
                    continue

                payload = {
                    "fields": {
                        "project": {"key": jira_key},
                        "summary": item.get("summary", "New Item"),
                        "issuetype": {"name": item.get("issuetype", "Task")},
                    }
                }
                desc = item.get("description") or ""
                if desc:
                    payload["fields"]["description"] = {
                        "type": "doc",
                        "version": 1,
                        "content": [{"type": "paragraph", "content": [{"type": "text", "text": desc}]}]
                    }
                priority = item.get("priority") or ""
                if priority:
                    payload["fields"]["priority"] = {"name": priority}
                if parent_key:
                    payload["fields"]["parent"] = {"key": parent_key}

                kwargs = {"headers": headers, "json": payload, "verify": JIRA_SSL_VERIFY, "timeout": 15}
                if auth:
                    kwargs["auth"] = auth

                r = requests.post(f"{JIRA_URL}/rest/api/3/issue", **kwargs)
                if r.status_code in (200, 201):
                    key = r.json().get("key")
                    resolved_jira_key[item["id"]] = key
                    synced_updates.append({"id": item["id"], "jira_key": key})
                    results.append({"id": item["id"], "success": True, "jira_key": key})
                else:
                    results.append({"id": item["id"], "success": False, "error": r.text})
            remaining = still_waiting

        for item in remaining:
            results.append({"id": item["id"], "success": False, "error": "Unresolved parent"})

        if synced_updates:
            try:
                db.mark_work_items_synced(project_id, synced_updates)
            except Exception:
                pass

        return JSONResponse({"success": True, "results": results})
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)


@app.delete("/jira/work_items/{item_id}")
async def delete_work_item(item_id: int):
    """
    Deletes a work item from the local tree, and (best-effort) the matching
    Jira issue if it was already synced - including any of its descendants
    that were synced too (e.g. deleting a Story also removes its Subtasks in
    Jira). Local removal always cascades via the FK regardless of whether
    the Jira-side deletes all succeeded, so a Jira hiccup never blocks
    clearing something out of the planning view.
    """
    subtree = db.get_work_item_subtree(item_id)
    if not subtree:
        return JSONResponse({"error": "Work item not found"}, status_code=404)

    jira_deleted = []
    jira_failed = []

    synced_items = [i for i in subtree if i.get("jira_key")]
    if synced_items:
        if not jira_configured():
            return JSONResponse({"error": "Jira not configured"}, status_code=503)
        headers, auth = jira_headers()
        for item in synced_items:
            kwargs = {"headers": headers, "params": {"deleteSubtasks": "true"}, "verify": JIRA_SSL_VERIFY, "timeout": 15}
            if auth:
                kwargs["auth"] = auth
            r = requests.delete(f"{JIRA_URL}/rest/api/3/issue/{item['jira_key']}", **kwargs)
            if r.status_code in (204, 404):
                jira_deleted.append(item["jira_key"])
            else:
                jira_failed.append({"key": item["jira_key"], "error": r.text})

    db.delete_work_item(item_id)

    return JSONResponse({"success": True, "jira_deleted": jira_deleted, "jira_failed": jira_failed})


@app.delete("/jira/issues/{issue_key}")
async def delete_jira_issue(issue_key: str):
    """
    Deletes a Jira issue directly by key - used from the 'Live in Jira' pane,
    which may show issues that were never tracked as a local work item (e.g.
    created before this feature, or outside this app). If a local work item
    does point at this key, it reverts to an unsynced draft instead of being
    left with a dangling reference.
    """
    if not jira_configured():
        return JSONResponse({"error": "Jira not configured"}, status_code=503)
    try:
        headers, auth = jira_headers()
        kwargs = {"headers": headers, "params": {"deleteSubtasks": "true"}, "verify": JIRA_SSL_VERIFY, "timeout": 15}
        if auth:
            kwargs["auth"] = auth
        r = requests.delete(f"{JIRA_URL}/rest/api/3/issue/{issue_key}", **kwargs)
        if r.status_code in (204, 404):
            try:
                db.clear_work_item_sync_by_jira_key(issue_key)
            except Exception:
                pass
            return JSONResponse({"success": True})
        return JSONResponse({"success": False, "error": r.text})
    except Exception as e:
        return JSONResponse({"success": False, "error": str(e)})


# ============================================================
# CHANGE LOG (Change Request Register)
# ============================================================

CR_TYPES = ("Scope Change", "Enhancement", "Defect Fix", "Compliance")
CR_STAGES = ("During Build", "Post Go-Live")
CR_STATUSES = ("Proposed", "Approved", "Rejected", "Deferred", "Implemented")


@app.post("/change_requests")
async def create_change_request(request: Request):
    body = await request.json()
    project_id = body.get("project_id")
    title = (body.get("title") or "").strip()
    if not project_id or not title:
        return JSONResponse({"error": "project_id and title are required"}, status_code=400)

    fields = {
        "title": title,
        "type": body.get("type") if body.get("type") in CR_TYPES else None,
        "stage": body.get("stage") if body.get("stage") in CR_STAGES else None,
        "priority": body.get("priority") or "Medium",
        "description": (body.get("description") or "").strip(),
        "impact_summary": (body.get("impact_summary") or "").strip(),
        "work_item_code": body.get("work_item_code") or None,
        "raised_by": (body.get("raised_by") or "").strip(),
    }
    try:
        cr = db.create_change_request(project_id, fields)
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)
    return JSONResponse({"success": True, "change_request": cr})


@app.get("/change_requests/{project_id}")
async def get_change_requests(project_id: str):
    try:
        rows = db.list_change_requests(project_id)
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)
    return JSONResponse({"change_requests": rows})


@app.put("/change_requests/{cr_id}")
async def edit_change_request(cr_id: int, request: Request):
    body = await request.json()
    fields = {}
    for col in ("title", "description", "impact_summary", "raised_by", "approved_by", "decision_at", "work_item_code"):
        if col in body:
            fields[col] = body[col]
    if "type" in body and (body["type"] in CR_TYPES or body["type"] is None):
        fields["type"] = body["type"]
    if "stage" in body and (body["stage"] in CR_STAGES or body["stage"] is None):
        fields["stage"] = body["stage"]
    if "priority" in body:
        fields["priority"] = body["priority"]
    if "status" in body and body["status"] in CR_STATUSES:
        fields["status"] = body["status"]

    cr = db.update_change_request(cr_id, fields)
    if not cr:
        return JSONResponse({"error": "Change request not found"}, status_code=404)
    return JSONResponse({"success": True, "change_request": cr})


@app.delete("/change_requests/{cr_id}")
async def remove_change_request(cr_id: int):
    db.delete_change_request(cr_id)
    return JSONResponse({"success": True})


# ============================================================
# CONFLUENCE INTEGRATION
# ============================================================

@app.post("/confluence/publish")
async def confluence_publish(request: Request):
    """
    Publishes a generated document directly to Confluence as a wiki page.
    If `content` is provided in the request body, that is published as-is
    (used when the caller drags a specific artifact over). Otherwise falls
    back to the BRD on disk for backward compatibility.
    """
    if not jira_configured():
        return JSONResponse({"error": "Atlassian credentials not configured"}, status_code=503)
    try:
        body = await request.json()
        title = body.get("title", "Banking Deposit page BRD")
        content = body.get("content")

        if content is None:
            project_dir, _ = get_project_dir(body.get("project_id"))
            brd_path = os.path.join(project_dir, "BRD.md")
            if not os.path.exists(brd_path):
                # Legacy fallback: projects created before per-project folders existed
                brd_path = os.path.join(WORKSPACE_DIR, "BRD.md")
            if os.path.exists(brd_path):
                with open(brd_path, "r", encoding="utf-8") as f:
                    content = f.read()
            else:
                content = "Requirements details not finalized."

        # Parse simple markdown to Confluence Storage HTML format
        html_content = content.replace("\n", "<br>").replace("### ", "<h3>").replace("## ", "<h2>").replace("# ", "<h1>")

        headers, auth = jira_headers()
        _, space_key = resolve_project_keys(body.get("project_id"))

        space_id = get_confluence_space_id(space_key, headers, auth)
        if not space_id:
            return JSONResponse({"success": False, "error": f"No Confluence space found with key '{space_key}'. Create a space with that key first."})

        payload = {
            "spaceId": space_id,
            "status": "current",
            "title": title,
            "body": {
                "representation": "storage",
                "value": f"<p>{html_content}</p>"
            }
        }

        kwargs = {"headers": headers, "json": payload, "verify": JIRA_SSL_VERIFY, "timeout": 15}
        if auth:
            kwargs["auth"] = auth

        # 2. Create the page via confluence REST v2 API
        r = requests.post(f"{JIRA_URL}/wiki/api/v2/pages", **kwargs)
        if r.status_code in (200, 201):
            page_id = r.json().get("id")
            page_url = f"{JIRA_URL}/wiki/spaces/{space_key}/pages/{page_id}"
            if body.get("project_id"):
                try:
                    db.add_confluence_synced(body.get("project_id"), title, page_url)
                except Exception:
                    pass
            return JSONResponse({"success": True, "url": page_url})
        else:
            return JSONResponse({"success": False, "error": r.text})
    except Exception as e:
        return JSONResponse({"success": False, "error": str(e)})


@app.get("/confluence/pages")
async def confluence_pages(project_id: str = None):
    """
    Live query of what pages actually exist right now in the project's
    Confluence space (not a locally-tracked record of what this app has
    pushed) — used to show the real current state in the Artifacts tab.
    """
    if not jira_configured():
        return JSONResponse({"error": "Atlassian credentials not configured"}, status_code=503)
    try:
        headers, auth = jira_headers()
        _, space_key = resolve_project_keys(project_id)

        space_id = get_confluence_space_id(space_key, headers, auth)
        if not space_id:
            return JSONResponse({"success": True, "pages": []})

        pages_kwargs = {"headers": headers, "params": {"space-id": space_id, "limit": 50, "sort": "-modified-date"}, "verify": JIRA_SSL_VERIFY, "timeout": 15}
        if auth:
            pages_kwargs["auth"] = auth
        pages_r = requests.get(f"{JIRA_URL}/wiki/api/v2/pages", **pages_kwargs)
        if pages_r.status_code != 200:
            return JSONResponse({"success": False, "error": pages_r.text})

        pages = [
            {
                "id": p.get("id"),
                "title": p.get("title"),
                "url": f"{JIRA_URL}/wiki/spaces/{space_key}/pages/{p.get('id')}"
            }
            for p in pages_r.json().get("results", [])
        ]
        return JSONResponse({"success": True, "pages": pages, "space_key": space_key})
    except Exception as e:
        return JSONResponse({"success": False, "error": str(e)})


@app.delete("/confluence/pages/{page_id}")
async def delete_confluence_page(page_id: str):
    if not jira_configured():
        return JSONResponse({"error": "Atlassian credentials not configured"}, status_code=503)
    try:
        headers, auth = jira_headers()
        kwargs = {"headers": headers, "verify": JIRA_SSL_VERIFY, "timeout": 15}
        if auth:
            kwargs["auth"] = auth
        r = requests.delete(f"{JIRA_URL}/wiki/api/v2/pages/{page_id}", **kwargs)
        if r.status_code in (200, 204):
            return JSONResponse({"success": True})
        return JSONResponse({"success": False, "error": r.text})
    except Exception as e:
        return JSONResponse({"success": False, "error": str(e)})


# ============================================================
# CORE GET ROUTES
# ============================================================

@app.get("/", response_class=HTMLResponse)
async def get_index():
    with open("index.html", "r", encoding="utf-8") as f:
        return f.read()

def _scan_artifacts_dir(dir_path, url_prefix):
    """Only .docx specs are recoverable - HTML diagram pages are intentionally
    not surfaced (see /chat/generate_artifacts, which stopped generating them)."""
    files_list = []
    if os.path.isdir(dir_path):
        for fname in sorted(os.listdir(dir_path)):
            base, ext = os.path.splitext(fname)
            if ext == ".docx":
                entry = {"name": fname, "url": f"{url_prefix}/{fname}"}
                md_path = os.path.join(dir_path, base + ".md")
                if os.path.exists(md_path):
                    with open(md_path, "r", encoding="utf-8") as f:
                        entry["content"] = f.read()
                files_list.append(entry)
    return files_list

@app.get("/artifacts/recover")
async def artifacts_recover(project_id: str = None):
    """
    Rebuilds an artifact manifest from whatever .docx/.html files already
    exist on disk for this project. Used to recover a project's Artifacts
    tab when the browser's saved manifest was lost (e.g. generated before
    client-side persistence was added). Reads the matching .md sidecar
    (same base filename) as the document's viewable/draggable content when
    present. Miro board links cannot be recovered this way since they are
    never written to disk. Falls back to the legacy workspace root if the
    project's own subfolder is empty (projects created before per-project
    folders existed).
    """
    project_dir, safe_project_id = get_project_dir(project_id)
    files_list = _scan_artifacts_dir(project_dir, f"/workspace/{safe_project_id}")
    if not files_list:
        files_list = _scan_artifacts_dir(WORKSPACE_DIR, "/workspace")
        if project_id and files_list:
            try:
                db.add_artifacts(project_id, files_list)
            except Exception:
                pass
    return JSONResponse({"success": True, "files": files_list})

@app.delete("/artifacts/file/{filename}")
async def delete_artifact_file(filename: str, project_id: str = None):
    """
    Deletes a generated artifact (and its .md sidecar, if any). Tries the
    project's own subfolder first, then falls back to the legacy workspace
    root for projects created before per-project folders existed.
    """
    safe_name = os.path.basename(filename)
    if not safe_name or safe_name in (".", ".."):
        return JSONResponse({"success": False, "error": "Invalid filename"}, status_code=400)

    project_dir, _ = get_project_dir(project_id)
    file_path = None
    for base_dir in (project_dir, WORKSPACE_DIR):
        candidate = os.path.join(base_dir, safe_name)
        if os.path.isfile(candidate):
            file_path = candidate
            break

    if not file_path:
        return JSONResponse({"success": False, "error": "File not found"}, status_code=404)

    try:
        os.remove(file_path)
        base, _ = os.path.splitext(file_path)
        md_path = base + ".md"
        if os.path.exists(md_path):
            os.remove(md_path)
        if project_id:
            try:
                db.delete_artifact_by_name(project_id, safe_name)
            except Exception:
                pass
        return JSONResponse({"success": True})
    except Exception as e:
        return JSONResponse({"success": False, "error": str(e)}, status_code=500)

@app.get("/jira/status")
async def jira_status(project_id: str = None):
    if not jira_configured():
        return JSONResponse({"connected": False, "error": "Not configured"})
    try:
        jira_key, _ = resolve_project_keys(project_id)
        headers, auth = jira_headers()
        kwargs = {"headers": headers, "verify": JIRA_SSL_VERIFY, "timeout": 8}
        if auth:
            kwargs["auth"] = auth
        r = requests.get(f"{JIRA_URL}/rest/api/3/myself", **kwargs)
        if r.status_code == 200:
            return JSONResponse({"connected": True, "project": jira_key, "jira_url": JIRA_URL})
        return JSONResponse({"connected": False})
    except Exception:
        return JSONResponse({"connected": False})

def _fetch_live_jira_issues(project_id, max_results=50):
    """Shared by /jira/issues and the Jira tasks CSV report - live JQL lookup
    of every issue in the project's Jira board, normalized to a flat dict."""
    jira_key, _ = resolve_project_keys(project_id)
    headers, auth = jira_headers()
    jql = f"project = {jira_key} ORDER BY updated DESC"
    params = {"jql": jql, "maxResults": max_results, "fields": "summary,status,assignee,priority,issuetype,parent,url"}
    kwargs = {"headers": headers, "params": params, "verify": JIRA_SSL_VERIFY, "timeout": 15}
    if auth:
        kwargs["auth"] = auth
    r = requests.get(f"{JIRA_URL}/rest/api/3/search/jql", **kwargs)
    r.raise_for_status()
    issues = []
    for issue in r.json().get("issues", []):
        f = issue.get("fields", {})
        status_obj = f.get("status", {})
        parent_obj = f.get("parent")
        issues.append({
            "key": issue["key"],
            "summary": f.get("summary", ""),
            "status": status_obj.get("name", "Unknown"),
            "priority": (f.get("priority") or {}).get("name", "Medium"),
            "issueType": (f.get("issuetype") or {}).get("name", "Task"),
            "parentKey": parent_obj.get("key") if parent_obj else None,
            "url": f"{JIRA_URL}/browse/{issue['key']}",
        })
    return issues


@app.get("/jira/issues")
async def jira_issues(status: str = "all", max_results: int = 50, project_id: str = None):
    if not jira_configured():
        return JSONResponse({"error": "Jira not configured"}, status_code=503)
    try:
        issues = _fetch_live_jira_issues(project_id, max_results)
        return JSONResponse({"issues": issues})
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)


def _get_rtm_rows_with_status(project_id):
    """Shared by /rtm/{project_id} and the RTM CSV report. Jira status is
    fetched live here (a single batched JQL lookup by key) rather than
    cached, since it changes outside this app; the lookup is best-effort so
    a Jira hiccup still returns the local matrix with jira_status left null."""
    rows = db.get_rtm_rows(project_id)

    jira_keys = sorted({r["jira_key"] for r in rows if r.get("jira_key")} | {r["test_case_jira_key"] for r in rows if r.get("test_case_jira_key")})
    status_by_key = {}
    url_by_key = {}
    if jira_keys and jira_configured():
        try:
            headers, auth = jira_headers()
            jql = "key in (" + ",".join(jira_keys) + ")"
            params = {"jql": jql, "maxResults": len(jira_keys), "fields": "status"}
            kwargs = {"headers": headers, "params": params, "verify": JIRA_SSL_VERIFY, "timeout": 15}
            if auth:
                kwargs["auth"] = auth
            r = requests.get(f"{JIRA_URL}/rest/api/3/search/jql", **kwargs)
            if r.status_code == 200:
                for issue in r.json().get("issues", []):
                    status_by_key[issue["key"]] = (issue.get("fields", {}).get("status") or {}).get("name", "Unknown")
                    url_by_key[issue["key"]] = f"{JIRA_URL}/browse/{issue['key']}"
        except Exception:
            pass

    for row in rows:
        key = row.get("jira_key")
        row["jira_status"] = status_by_key.get(key) if key else None
        row["jira_url"] = url_by_key.get(key) if key else None
        tc_key = row.get("test_case_jira_key")
        row["test_case_jira_status"] = status_by_key.get(tc_key) if tc_key else None
        row["test_case_jira_url"] = url_by_key.get(tc_key) if tc_key else None
    return rows


@app.get("/rtm/{project_id}")
async def get_rtm(project_id: str):
    """
    Requirements Traceability Matrix: Epic -> Story/Task -> Test Case, one
    row per (requirement, test case) pair - requirements with no test case
    yet get a single "not covered" row.
    """
    try:
        rows = _get_rtm_rows_with_status(project_id)
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)

    item_ids = {r["item_id"] for r in rows}
    covered_ids = {r["item_id"] for r in rows if r["covered"]}
    total = len(item_ids)
    covered = len(covered_ids)

    try:
        test_cases = db.get_test_cases(project_id)
    except Exception:
        test_cases = []
    test_pass = sum(1 for tc in test_cases if tc.get("status") == "Pass")
    test_fail = sum(1 for tc in test_cases if tc.get("status") == "Fail")
    test_executed = test_pass + test_fail

    return JSONResponse({
        "rows": rows,
        "summary": {
            "total_requirements": total,
            "covered": covered,
            "not_covered": total - covered,
            "coverage_pct": round((covered / total) * 100, 1) if total else 0,
            "test_pass": test_pass,
            "test_fail": test_fail,
            "test_executed": test_executed,
            "test_pass_rate_pct": round((test_pass / test_executed) * 100, 1) if test_executed else 0,
        }
    })


def _csv_response(fieldnames, rows, filename):
    buf = io.StringIO()
    writer = csv.DictWriter(buf, fieldnames=fieldnames, extrasaction="ignore")
    writer.writeheader()
    for row in rows:
        writer.writerow(row)
    return Response(content=buf.getvalue(), media_type="text/csv", headers={"Content-Disposition": f'attachment; filename="{filename}"'})


def _jira_tasks_report_data(project_id):
    """Shared by the CSV download and the Confluence sync - same fieldnames/
    rows either way so the two outputs never drift apart."""
    issues = _fetch_live_jira_issues(project_id, max_results=200)
    fieldnames = ["Key", "Summary", "Type", "Status", "Priority", "Parent", "URL"]
    rows = [{
        "Key": i["key"], "Summary": i["summary"], "Type": i["issueType"], "Status": i["status"],
        "Priority": i["priority"], "Parent": i.get("parentKey") or "", "URL": i["url"],
    } for i in issues]
    return fieldnames, rows


def _rtm_report_data(project_id):
    rows = _get_rtm_rows_with_status(project_id)
    fieldnames = ["Epic", "Epic Summary", "Requirement", "Type", "Requirement Summary", "Priority",
                  "Jira Key", "Jira Status", "Test Case", "Test Case Title", "Test Status",
                  "Test Jira Key", "Test Jira Status", "Covered"]
    out = [{
        "Epic": r.get("epic_code") or "", "Epic Summary": r.get("epic_summary") or "",
        "Requirement": r.get("item_code") or "", "Type": r.get("item_type") or "",
        "Requirement Summary": r.get("item_summary") or "", "Priority": r.get("priority") or "",
        "Jira Key": r.get("jira_key") or "", "Jira Status": r.get("jira_status") or "",
        "Test Case": r.get("test_case_code") or "", "Test Case Title": r.get("test_case_title") or "",
        "Test Status": r.get("test_case_status") or "", "Test Jira Key": r.get("test_case_jira_key") or "",
        "Test Jira Status": r.get("test_case_jira_status") or "", "Covered": "Yes" if r.get("covered") else "No",
    } for r in rows]
    return fieldnames, out


def _change_log_report_data(project_id):
    rows = db.list_change_requests(project_id)
    fieldnames = ["Code", "Title", "Type", "Stage", "Priority", "Status", "Linked Requirement",
                  "Description", "Impact Summary", "Raised By", "Raised At", "Approved By", "Decision At"]
    out = [{
        "Code": r.get("code") or "", "Title": r.get("title") or "", "Type": r.get("type") or "",
        "Stage": r.get("stage") or "", "Priority": r.get("priority") or "", "Status": r.get("status") or "",
        "Linked Requirement": r.get("work_item_code") or "", "Description": r.get("description") or "",
        "Impact Summary": r.get("impact_summary") or "", "Raised By": r.get("raised_by") or "",
        "Raised At": r.get("raised_at") or "", "Approved By": r.get("approved_by") or "",
        "Decision At": r.get("decision_at") or "",
    } for r in rows]
    return fieldnames, out


REPORT_BUILDERS = {
    "jira_tasks": ("Live Jira Tasks", "jira_live_tasks.csv", _jira_tasks_report_data),
    "rtm": ("RTM / Trace Matrix", "rtm.csv", _rtm_report_data),
    "change_log": ("Change Log", "change_log.csv", _change_log_report_data),
}


@app.get("/reports/jira_tasks.csv")
async def report_jira_tasks_csv(project_id: str = None):
    if not jira_configured():
        return JSONResponse({"error": "Jira not configured"}, status_code=503)
    try:
        fieldnames, rows = _jira_tasks_report_data(project_id)
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)
    return _csv_response(fieldnames, rows, "jira_live_tasks.csv")


@app.get("/reports/rtm.csv")
async def report_rtm_csv(project_id: str):
    try:
        fieldnames, rows = _rtm_report_data(project_id)
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)
    return _csv_response(fieldnames, rows, "rtm.csv")


@app.get("/reports/change_log.csv")
async def report_change_log_csv(project_id: str):
    try:
        fieldnames, rows = _change_log_report_data(project_id)
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)
    return _csv_response(fieldnames, rows, "change_log.csv")


def _esc_html(s):
    return str(s if s is not None else "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _rows_to_html_table(fieldnames, rows):
    thead = "".join(f"<th>{_esc_html(f)}</th>" for f in fieldnames)
    body_rows = "".join(
        "<tr>" + "".join(f"<td>{_esc_html(row.get(f, ''))}</td>" for f in fieldnames) + "</tr>"
        for row in rows
    ) or f'<tr><td colspan="{len(fieldnames)}">No data yet</td></tr>'
    return f"<table><thead><tr>{thead}</tr></thead><tbody>{body_rows}</tbody></table>"


def _confluence_upsert_page(project_id, title, html_body):
    """
    Creates a Confluence page, or REPLACES it in place if a page with this
    exact title already exists in the project's space - so re-syncing a
    report updates the same page (new version) instead of piling up
    duplicates every time it's re-synced.
    """
    headers, auth = jira_headers()
    _, space_key = resolve_project_keys(project_id)
    space_id = get_confluence_space_id(space_key, headers, auth)
    if not space_id:
        return None, f"No Confluence space found with key '{space_key}'. Create a space with that key first."

    get_kwargs = {"headers": headers, "params": {"space-id": space_id, "title": title}, "verify": JIRA_SSL_VERIFY, "timeout": 15}
    if auth:
        get_kwargs["auth"] = auth
    find_r = requests.get(f"{JIRA_URL}/wiki/api/v2/pages", **get_kwargs)
    existing_id = None
    if find_r.status_code == 200:
        results = find_r.json().get("results", [])
        if results:
            existing_id = results[0]["id"]

    body = {"representation": "storage", "value": html_body}
    if existing_id:
        version_kwargs = {"headers": headers, "verify": JIRA_SSL_VERIFY, "timeout": 15}
        if auth:
            version_kwargs["auth"] = auth
        version_r = requests.get(f"{JIRA_URL}/wiki/api/v2/pages/{existing_id}", **version_kwargs)
        if version_r.status_code != 200:
            return None, version_r.text
        current_version = (version_r.json().get("version") or {}).get("number", 1)
        payload = {"id": existing_id, "status": "current", "title": title, "body": body, "version": {"number": current_version + 1}}
        put_kwargs = {"headers": headers, "json": payload, "verify": JIRA_SSL_VERIFY, "timeout": 15}
        if auth:
            put_kwargs["auth"] = auth
        r = requests.put(f"{JIRA_URL}/wiki/api/v2/pages/{existing_id}", **put_kwargs)
        if r.status_code in (200, 201):
            return f"{JIRA_URL}/wiki/spaces/{space_key}/pages/{existing_id}", None
        return None, r.text
    else:
        payload = {"spaceId": space_id, "status": "current", "title": title, "body": body}
        post_kwargs = {"headers": headers, "json": payload, "verify": JIRA_SSL_VERIFY, "timeout": 15}
        if auth:
            post_kwargs["auth"] = auth
        r = requests.post(f"{JIRA_URL}/wiki/api/v2/pages", **post_kwargs)
        if r.status_code in (200, 201):
            page_id = r.json().get("id")
            return f"{JIRA_URL}/wiki/spaces/{space_key}/pages/{page_id}", None
        return None, r.text


@app.post("/reports/sync_to_confluence")
async def sync_report_to_confluence(request: Request):
    if not jira_configured():
        return JSONResponse({"error": "Atlassian credentials not configured"}, status_code=503)
    body = await request.json()
    project_id = body.get("project_id")
    report = body.get("report")
    if not project_id or report not in REPORT_BUILDERS:
        return JSONResponse({"error": "project_id and a valid report are required"}, status_code=400)

    report_label, _, builder = REPORT_BUILDERS[report]
    try:
        fieldnames, rows = builder(project_id)
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)

    proj_basic = db.get_project_basic(project_id)
    project_name = (proj_basic or {}).get("name") or "Project"
    title = f"{project_name} - {report_label}"
    html_table = _rows_to_html_table(fieldnames, rows)

    url, err = _confluence_upsert_page(project_id, title, html_table)
    if err:
        return JSONResponse({"error": err}, status_code=502)
    try:
        db.add_confluence_synced(project_id, title, url)
    except Exception:
        pass
    return JSONResponse({"success": True, "url": url, "title": title})


ISSUETYPE_ORDER = ["Epic", "Story", "Task", "Subtask"]
PRIORITY_ORDER = ["High", "Medium", "Low"]
TEST_STATUS_ORDER = ["Pass", "Fail", "Not Run"]


def _ordered_counts(counts, order):
    """Sorts a {label: count} dict into a canonical display order, with any
    unexpected labels appended after (alphabetically) rather than dropped."""
    known = [{"label": k, "count": counts[k]} for k in order if k in counts]
    extra = sorted(k for k in counts if k not in order)
    known += [{"label": k, "count": counts[k]} for k in extra]
    return known


@app.get("/dashboard/{project_id}")
async def get_dashboard(project_id: str):
    """
    Aggregates everything the project dashboard needs into one call: work
    item mix, Jira sync rate, RTM coverage, test case status, live Jira
    status (best-effort), and per-Epic sync progress. Each section degrades
    independently on failure (empty/zero) rather than failing the whole
    dashboard over one flaky piece (matches the resilience pattern used by
    /rtm and /jira/issues elsewhere in this file).
    """
    try:
        items = db.get_all_work_items(project_id)
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)

    type_counts, priority_counts = {}, {}
    synced = 0
    for it in items:
        t = it.get("issuetype", "Unknown")
        type_counts[t] = type_counts.get(t, 0) + 1
        if it.get("jira_key"):
            synced += 1
        if it.get("issuetype") in ("Epic", "Story", "Task"):
            p = it.get("priority") or "Unassigned"
            priority_counts[p] = priority_counts.get(p, 0) + 1

    total_items = len(items)
    sync_rate = round((synced / total_items) * 100, 1) if total_items else 0.0

    try:
        rtm_rows = db.get_rtm_rows(project_id)
        req_ids = {r["item_id"] for r in rtm_rows}
        covered_ids = {r["item_id"] for r in rtm_rows if r["covered"]}
        total_reqs, covered = len(req_ids), len(covered_ids)
    except Exception:
        total_reqs, covered = 0, 0
    coverage_pct = round((covered / total_reqs) * 100, 1) if total_reqs else 0.0

    try:
        test_cases = db.get_test_cases(project_id)
    except Exception:
        test_cases = []
    tc_status_counts = {}
    for tc in test_cases:
        s = tc.get("status") or "Not Run"
        tc_status_counts[s] = tc_status_counts.get(s, 0) + 1
    test_pass = tc_status_counts.get("Pass", 0)
    test_fail = tc_status_counts.get("Fail", 0)
    test_executed = test_pass + test_fail
    test_pass_rate_pct = round((test_pass / test_executed) * 100, 1) if test_executed else 0.0

    try:
        change_requests = db.list_change_requests(project_id)
    except Exception:
        change_requests = []
    cr_stage_counts, cr_status_counts = {}, {}
    for cr in change_requests:
        cr_stage_counts[cr.get("stage") or "Unspecified"] = cr_stage_counts.get(cr.get("stage") or "Unspecified", 0) + 1
        cr_status_counts[cr.get("status") or "Proposed"] = cr_status_counts.get(cr.get("status") or "Proposed", 0) + 1
    cr_total = len(change_requests)
    cr_during_build = cr_stage_counts.get("During Build", 0)
    cr_post_live = cr_stage_counts.get("Post Go-Live", 0)
    cr_pending = cr_status_counts.get("Proposed", 0)

    try:
        proj = db.get_project(project_id)
    except Exception:
        proj = None
    artifacts_count = len(proj.get("artifacts", [])) if proj else 0
    confluence_count = len(proj.get("confluence_synced", [])) if proj else 0

    jira_status_counts = {}
    jira_open = jira_done = 0
    if jira_configured():
        try:
            jira_key, _ = resolve_project_keys(project_id)
            headers, auth = jira_headers()
            params = {"jql": f"project = {jira_key} ORDER BY updated DESC", "maxResults": 100, "fields": "status"}
            kwargs = {"headers": headers, "params": params, "verify": JIRA_SSL_VERIFY, "timeout": 15}
            if auth:
                kwargs["auth"] = auth
            r = requests.get(f"{JIRA_URL}/rest/api/3/search/jql", **kwargs)
            if r.status_code == 200:
                for issue in r.json().get("issues", []):
                    st = (issue.get("fields", {}).get("status") or {}).get("name", "Unknown")
                    jira_status_counts[st] = jira_status_counts.get(st, 0) + 1
                    if any(k in st.lower() for k in ("done", "closed", "resolved")):
                        jira_done += 1
                    else:
                        jira_open += 1
        except Exception:
            pass

    by_id = {it["id"]: it for it in items}
    children_of = {}
    for it in items:
        if it.get("parent_id"):
            children_of.setdefault(it["parent_id"], []).append(it["id"])

    def collect_subtree(node_id):
        acc = [node_id]
        for c in children_of.get(node_id, []):
            acc.extend(collect_subtree(c))
        return acc

    epic_progress = []
    for it in items:
        if it.get("issuetype") == "Epic":
            subtree_items = [by_id[i] for i in collect_subtree(it["id"])]
            total = len(subtree_items)
            synced_n = sum(1 for x in subtree_items if x.get("jira_key"))
            epic_progress.append({
                "code": it.get("code"),
                "summary": it.get("summary"),
                "total": total,
                "synced": synced_n,
                "pct": round((synced_n / total) * 100, 1) if total else 0.0
            })
    epic_progress.sort(key=lambda e: e.get("code") or "")

    return JSONResponse({
        "kpis": {
            "total_work_items": total_items,
            "total_requirements": total_reqs,
            "jira_synced": synced,
            "jira_total": total_items,
            "sync_rate_pct": sync_rate,
            "coverage_pct": coverage_pct,
            "covered": covered,
            "not_covered": total_reqs - covered,
            "test_executed": test_executed,
            "test_total": len(test_cases),
            "test_pass": test_pass,
            "test_fail": test_fail,
            "test_not_run": tc_status_counts.get("Not Run", 0),
            "test_pass_rate_pct": test_pass_rate_pct,
            "artifacts_count": artifacts_count,
            "confluence_synced_count": confluence_count,
            "jira_open": jira_open,
            "jira_done": jira_done,
            "cr_total": cr_total,
            "cr_during_build": cr_during_build,
            "cr_post_live": cr_post_live,
            "cr_pending": cr_pending,
        },
        "work_item_breakdown": _ordered_counts(type_counts, ISSUETYPE_ORDER),
        "priority_distribution": _ordered_counts(priority_counts, PRIORITY_ORDER),
        "test_case_status": _ordered_counts(tc_status_counts, TEST_STATUS_ORDER),
        "jira_status_breakdown": sorted(
            ({"label": k, "count": v} for k, v in jira_status_counts.items()),
            key=lambda x: -x["count"]
        ),
        "change_request_stage_breakdown": _ordered_counts(cr_stage_counts, CR_STAGES),
        "change_request_status_breakdown": _ordered_counts(cr_status_counts, CR_STATUSES),
        "epic_progress": epic_progress
    })


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8003)
